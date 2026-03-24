import copy
import shutil
import tempfile
import unittest
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from thesis_config import DEFAULT_CONFIG
from thesis_formatter import numbering
from thesis_formatter.formatter import apply_format


PROJECT_ROOT = Path(__file__).resolve().parents[1]
ARTIFACT_ROOT = PROJECT_ROOT / "tmp_test_artifacts" / "caption_eval_20260320_rerun"
IMAGE_PATH = PROJECT_ROOT / "image.png"


def _document_texts(docx_path):
    doc = Document(str(docx_path))
    return [para.text.strip() for para in doc.paragraphs if para.text.strip()]


def _document_xml(docx_path):
    with zipfile.ZipFile(docx_path, "r") as zf:
        return zf.read("word/document.xml").decode("utf-8", errors="ignore")




def _normal_style_xml(docx_path):
    with zipfile.ZipFile(docx_path, "r") as zf:
        root = ET.fromstring(zf.read("word/styles.xml"))
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    style = root.find("w:style[@w:styleId='Normal']", ns)
    return ET.tostring(style, encoding="unicode") if style is not None else ""


def _style_spacing_attrs(docx_path, style_id):
    with zipfile.ZipFile(docx_path, "r") as zf:
        root = ET.fromstring(zf.read("word/styles.xml"))
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    style = root.find(f"w:style[@w:styleId='{style_id}']", ns)
    if style is None:
        return {}
    spacing = style.find("w:pPr/w:spacing", ns)
    if spacing is None:
        return {}
    return {
        "line": spacing.get(qn("w:line")),
        "before": spacing.get(qn("w:before")),
        "beforeLines": spacing.get(qn("w:beforeLines")),
        "after": spacing.get(qn("w:after")),
        "afterLines": spacing.get(qn("w:afterLines")),
    }


def _paragraph_xml(docx_path, index, strip_sectpr=False):
    doc = Document(str(docx_path))
    para_el = copy.deepcopy(doc.paragraphs[index]._element)
    if strip_sectpr:
        ppr = para_el.find(qn("w:pPr"))
        if ppr is not None:
            for sectpr in list(ppr.findall(qn("w:sectPr"))):
                ppr.remove(sectpr)
            if len(ppr) == 0:
                para_el.remove(ppr)
    normalized = ET.fromstring(para_el.xml.encode("utf-8"))
    return ET.tostring(normalized, encoding="unicode")

def _strip_numbering_part(docx_path):
    with zipfile.ZipFile(docx_path, "r") as zin:
        entries = {info.filename: zin.read(info.filename) for info in zin.infolist()}

    entries.pop("word/numbering.xml", None)

    rels_root = ET.fromstring(entries["word/_rels/document.xml.rels"])
    rels_ns = {"pr": "http://schemas.openxmlformats.org/package/2006/relationships"}
    for rel in list(rels_root.findall("pr:Relationship", rels_ns)):
        if rel.get("Type", "").endswith("/numbering"):
            rels_root.remove(rel)
    entries["word/_rels/document.xml.rels"] = ET.tostring(rels_root, encoding="utf-8", xml_declaration=True)

    ct_root = ET.fromstring(entries["[Content_Types].xml"])
    ct_ns = {"ct": "http://schemas.openxmlformats.org/package/2006/content-types"}
    for override in list(ct_root.findall("ct:Override", ct_ns)):
        if override.get("PartName") == "/word/numbering.xml":
            ct_root.remove(override)
    entries["[Content_Types].xml"] = ET.tostring(ct_root, encoding="utf-8", xml_declaration=True)

    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for filename, data in entries.items():
            zout.writestr(filename, data)


class CaptionModeTests(unittest.TestCase):
    def setUp(self):
        self.tmpdir = Path(tempfile.mkdtemp(prefix="caption_mode_test_", dir=PROJECT_ROOT / "tmp_test_artifacts"))

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def _run_apply(self, input_path, cfg=None, output_name=None):
        cfg = copy.deepcopy(cfg or DEFAULT_CONFIG)
        output_path = self.tmpdir / (output_name or f"{Path(input_path).stem}_output.docx")
        warnings = apply_format(str(input_path), str(output_path), config=cfg) or []
        return output_path, warnings, cfg

    def _base_dynamic_cfg(self):
        cfg = copy.deepcopy(DEFAULT_CONFIG)
        cfg["captions"]["mode"] = "dynamic"
        cfg["cover"]["enabled"] = False
        cfg["captions"]["check_numbering"] = False
        return cfg

    def test_stable_regression_uses_existing_real_docx_inputs(self):
        cases = [
            ("basic_fig_input.docx", ["图1.1 系统结构图"], []),
            ("basic_table_input.docx", ["表1.1 实验结果"], []),
            ("continued_table_input.docx", ["表1.1 第一部分", "续表1.1 第二部分"], []),
            ("chapter_fig_input.docx", ["图1.1 第一章图题", "图2.1 第二章图题"], []),
            ("level2_fig_input.docx", ["图1.1.1 小节图题", "图1.2.1 第二小节图题"], []),
            ("gap_fig_input.docx", ["图1.1 第一张图", "图1.3 第三张图"], ["正文第1章图编号不连续"]),
        ]

        for filename, expected_texts, expected_warnings in cases:
            with self.subTest(filename=filename):
                cfg = copy.deepcopy(DEFAULT_CONFIG)
                cfg["captions"]["mode"] = "stable"
                output_path, warnings, runtime_cfg = self._run_apply(ARTIFACT_ROOT / filename, cfg=cfg)
                texts = _document_texts(output_path)
                for expected in expected_texts:
                    self.assertIn(expected, texts)
                for expected_warning in expected_warnings:
                    self.assertTrue(any(expected_warning in warning for warning in warnings), warnings)
                self.assertEqual(runtime_cfg.get("_runtime", {}).get("caption_mode_effective"), "stable")

    def test_dynamic_mode_falls_back_when_document_cannot_be_normalized(self):
        doc = Document()
        doc.add_paragraph("绪论")
        doc.add_paragraph("图1.1 无法标准化图题")
        input_path = self.tmpdir / "dynamic_unfixable_input.docx"
        doc.save(input_path)

        cfg = self._base_dynamic_cfg()
        cfg["sections"]["renumber_headings"] = False
        output_path, warnings, runtime_cfg = self._run_apply(
            input_path,
            cfg=cfg,
        )

        self.assertTrue(output_path.exists())
        self.assertEqual(runtime_cfg.get("_runtime", {}).get("caption_mode_effective"), "stable")
        self.assertFalse(runtime_cfg.get("_runtime", {}).get("caption_mode_raw_precheck_passed"))
        self.assertTrue(any("dynamic 题注模式预检未通过" in warning for warning in warnings), warnings)
        self.assertNotIn("STYLEREF", _document_xml(output_path))

    def test_dynamic_mode_can_enable_after_auto_heading_and_multilevel_normalization(self):
        doc = Document()
        doc.add_paragraph("第1章 绪论")
        doc.add_picture(str(IMAGE_PATH))
        doc.add_paragraph("图1.1 自动标准化后图题")
        input_path = self.tmpdir / "dynamic_auto_normalized_input.docx"
        doc.save(input_path)

        output_path, warnings, runtime_cfg = self._run_apply(
            input_path,
            cfg=self._base_dynamic_cfg(),
        )

        self.assertTrue(output_path.exists())
        self.assertEqual(runtime_cfg.get("_runtime", {}).get("caption_mode_effective"), "dynamic")
        self.assertFalse(runtime_cfg.get("_runtime", {}).get("caption_mode_raw_precheck_passed"))
        self.assertTrue(any("自动补 Heading/多级编号后已通过" in warning for warning in warnings), warnings)
        xml = _document_xml(output_path)
        self.assertIn("STYLEREF 1 \s", xml)
        self.assertIn("SEQ Figure \* ARABIC \s 1", xml)

    def test_dynamic_mode_passes_with_real_heading_numbering_and_writes_dynamic_fields(self):
        doc = Document()
        doc.add_paragraph("第1章 绪论", style="Heading 1")
        numbering.setup_multilevel_list(doc, copy.deepcopy(DEFAULT_CONFIG))
        doc.add_picture(str(IMAGE_PATH))
        doc.add_paragraph("图1.1 动态图题")
        input_path = self.tmpdir / "dynamic_valid_input.docx"
        doc.save(input_path)

        output_path, warnings, runtime_cfg = self._run_apply(
            input_path,
            cfg=self._base_dynamic_cfg(),
        )

        self.assertTrue(output_path.exists())
        self.assertEqual(runtime_cfg.get("_runtime", {}).get("caption_mode_effective"), "dynamic")
        self.assertTrue(any("dynamic 题注模式预检通过" in warning for warning in warnings), warnings)
        xml = _document_xml(output_path)
        self.assertIn("STYLEREF 1 \s", xml)
        self.assertIn("SEQ Figure \* ARABIC \s 1", xml)

    def test_missing_numbering_part_is_recreated_without_crashing(self):
        doc = Document()
        doc.add_paragraph("第1章 绪论", style="Heading 1")
        doc.add_picture(str(IMAGE_PATH))
        doc.add_paragraph("图1.1 缺失numbering part测试")
        input_path = self.tmpdir / "missing_numbering_part_input.docx"
        doc.save(input_path)
        _strip_numbering_part(input_path)

        output_path, warnings, runtime_cfg = self._run_apply(
            input_path,
            cfg=self._base_dynamic_cfg(),
        )

        self.assertTrue(output_path.exists())
        self.assertEqual(runtime_cfg.get("_runtime", {}).get("caption_mode_effective"), "dynamic")
        self.assertTrue(any("自动补 Heading/多级编号后已通过" in warning for warning in warnings), warnings)
        xml = _document_xml(output_path)
        self.assertIn("STYLEREF 1 \s", xml)

    def test_caption_position_and_continued_table_warnings_are_reported(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "2"
        doc.add_paragraph("表1.1 错位表题")
        doc.add_paragraph("续表1.2 编号错误")
        input_path = self.tmpdir / "caption_layout_warning_input.docx"
        doc.save(input_path)

        _, warnings, _ = self._run_apply(input_path)

        self.assertTrue(any("表题位置异常" in warning for warning in warnings), warnings)
        self.assertTrue(any("续表编号未延续上一表" in warning for warning in warnings), warnings)

    def test_front_matter_skip_preserves_existing_front_pages_and_normal_style(self):
        doc = Document()
        doc.add_paragraph("本科毕业论文(或设计)")
        doc.add_paragraph("原创性声明")
        doc.add_paragraph("摘要")
        doc.add_paragraph("这是摘要内容。")
        doc.add_paragraph("Abstract")
        doc.add_paragraph("This is abstract.")
        doc.add_paragraph("第1章 绪论", style="Heading 1")
        doc.add_paragraph("正文内容")
        input_path = self.tmpdir / "front_matter_skip_preserve_input.docx"
        doc.save(input_path)

        input_normal_style = _normal_style_xml(input_path)

        cfg = copy.deepcopy(DEFAULT_CONFIG)
        cfg["cover"]["enabled"] = False
        cfg["front_matter"]["mode"] = "skip"

        output_path, _, _ = self._run_apply(input_path, cfg=cfg)

        out_doc = Document(str(output_path))
        self.assertEqual(out_doc.paragraphs[0].text.strip(), "本科毕业论文(或设计)")
        self.assertEqual(out_doc.paragraphs[1].text.strip(), "原创性声明")
        self.assertEqual(out_doc.paragraphs[2].text.strip(), "摘要")
        self.assertEqual(_normal_style_xml(output_path), input_normal_style)

    def test_front_matter_skip_restores_original_paragraph_xml_in_stable_and_dynamic_modes(self):
        doc = Document()
        cover = doc.add_paragraph("本科毕业论文(或设计)")
        cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover.runs[0].font.size = Pt(26)

        declaration = doc.add_paragraph("原创性声明")
        declaration.alignment = WD_ALIGN_PARAGRAPH.CENTER
        declaration.runs[0].font.size = Pt(14)

        doc.add_paragraph("摘要", style="Heading 1")
        doc.add_paragraph("这是摘要内容。")
        doc.add_paragraph("关键词：测试；前置页")
        doc.add_paragraph("Abstract", style="Heading 1")
        doc.add_paragraph("Abstract: This is abstract.")
        doc.add_paragraph("Key words: test; front matter")
        doc.add_paragraph("第1章 绪论", style="Heading 1")
        doc.add_paragraph("正文内容")
        input_path = self.tmpdir / "front_matter_skip_restore_input.docx"
        doc.save(input_path)

        expected_xml = {
            idx: _paragraph_xml(input_path, idx, strip_sectpr=True)
            for idx in range(8)
        }

        for mode in ("stable", "dynamic"):
            with self.subTest(mode=mode):
                cfg = copy.deepcopy(DEFAULT_CONFIG)
                cfg["cover"]["enabled"] = False
                cfg["front_matter"]["mode"] = "skip"
                cfg["captions"]["mode"] = mode

                output_path, _, runtime_cfg = self._run_apply(
                    input_path,
                    cfg=cfg,
                    output_name=f"front_matter_skip_restore_{mode}.docx",
                )

                for idx, xml in expected_xml.items():
                    self.assertEqual(_paragraph_xml(output_path, idx, strip_sectpr=True), xml)
                self.assertEqual(runtime_cfg.get("_runtime", {}).get("caption_mode_requested"), mode)

    def test_existing_cover_is_preserved_when_cover_generation_is_disabled(self):
        doc = Document()
        cover_para = doc.add_paragraph("本科毕业论文(或设计)")
        cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_run = cover_para.runs[0]
        cover_run.font.size = Pt(26)

        title_para = doc.add_paragraph("我的论文题目")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.runs[0].font.size = Pt(22)

        doc.add_paragraph("摘要")
        doc.add_paragraph("这是摘要内容。")
        doc.add_paragraph("第1章 绪论", style="Heading 1")
        input_path = self.tmpdir / "existing_cover_preserved_input.docx"
        doc.save(input_path)

        cfg = copy.deepcopy(DEFAULT_CONFIG)
        cfg["cover"]["enabled"] = False
        cfg["front_matter"]["mode"] = "skip"

        output_path, _, _ = self._run_apply(input_path, cfg=cfg)

        out_doc = Document(str(output_path))
        out_cover = out_doc.paragraphs[0]
        self.assertEqual(out_cover.text.strip(), "本科毕业论文(或设计)")

    def test_toc_can_be_disabled_entirely(self):
        doc = Document()
        doc.add_paragraph("第1章 绪论", style="Heading 1")
        doc.add_paragraph("正文内容")
        input_path = self.tmpdir / "toc_disabled_input.docx"
        doc.save(input_path)

        cfg = copy.deepcopy(DEFAULT_CONFIG)
        cfg["cover"]["enabled"] = False
        cfg["toc"]["enabled"] = False

        output_path, _, _ = self._run_apply(input_path, cfg=cfg)

        xml = _document_xml(output_path)
        self.assertNotIn(' TOC \o "1-', xml)
        texts = _document_texts(output_path)
        self.assertNotIn("目        录", texts)

    def test_toc_styles_honor_configured_spacing(self):
        doc = Document()
        doc.add_paragraph("第1章 绪论", style="Heading 1")
        doc.add_paragraph("1.1 研究背景", style="Heading 2")
        input_path = self.tmpdir / "toc_spacing_input.docx"
        doc.save(input_path)

        cfg = copy.deepcopy(DEFAULT_CONFIG)
        cfg["cover"]["enabled"] = False
        cfg["toc"]["line_spacing"] = 2.0
        cfg["toc"]["space_before"] = 1.0
        cfg["toc"]["space_after"] = 0.5

        output_path, _, _ = self._run_apply(input_path, cfg=cfg)

        toc1_spacing = _style_spacing_attrs(output_path, "TOC1")
        toc2_spacing = _style_spacing_attrs(output_path, "TOC2")
        expected_line = str(int(cfg["toc"]["line_spacing"] * 240))
        expected_toc1 = {
            "line": expected_line,
            "before": None,
            "beforeLines": str(int(cfg["toc"]["space_before"] * 100)),
            "after": None,
            "afterLines": str(int(cfg["toc"]["space_after"] * 100)),
        }
        expected_toc2 = {
            "line": expected_line,
            "before": None,
            "beforeLines": str(int(cfg["toc"]["space_before"] * 100)),
            "after": None,
            "afterLines": str(int(cfg["toc"]["space_after"] * 100)),
        }
        self.assertEqual(toc1_spacing, expected_toc1)
        self.assertEqual(toc2_spacing, expected_toc2)

    def test_caption_line_spacing_can_differ_from_body(self):
        doc = Document()
        doc.add_paragraph("第1章 绪论", style="Heading 1")
        doc.add_picture(str(IMAGE_PATH))
        doc.add_paragraph("图1.1 自定义题注行距")
        input_path = self.tmpdir / "caption_line_spacing_input.docx"
        doc.save(input_path)

        cfg = copy.deepcopy(DEFAULT_CONFIG)
        cfg["cover"]["enabled"] = False
        cfg["body"]["line_spacing"] = 1.5
        cfg["captions"]["line_spacing"] = 2.0

        output_path, _, _ = self._run_apply(input_path, cfg=cfg)

        out_doc = Document(str(output_path))
        caption_para = next(para for para in out_doc.paragraphs if para.text.strip() == "图1.1 自定义题注行距")
        self.assertEqual(caption_para.paragraph_format.line_spacing, 2.0)



if __name__ == "__main__":
    unittest.main()
