import copy
import shutil
import tempfile
import unittest
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn

from thesis_config import DEFAULT_CONFIG
from thesis_formatter._common import (
    format_line_spacing_value,
    format_paragraph_spacing_value,
    line_spacing_to_ooxml,
    line_spacing_to_points_and_rule,
    normalize_line_spacing,
    normalize_paragraph_spacing,
)
from thesis_formatter.formatter import apply_format
from thesis_gui import FormatterGUI


PROJECT_ROOT = Path(__file__).resolve().parents[1]
IMAGE_PATH = PROJECT_ROOT / "image.png"
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _spacing_attrs(node):
    if node is None:
        return {}
    return {
        "line": node.get(qn("w:line")),
        "lineRule": node.get(qn("w:lineRule")),
        "before": node.get(qn("w:before")),
        "beforeLines": node.get(qn("w:beforeLines")),
        "after": node.get(qn("w:after")),
        "afterLines": node.get(qn("w:afterLines")),
    }


def _style_spacing_by_id(docx_path, style_id):
    with zipfile.ZipFile(docx_path, "r") as zf:
        root = ET.fromstring(zf.read("word/styles.xml"))
    style = root.find(f"w:style[@w:styleId='{style_id}']", NS)
    if style is None:
        return {}
    return _spacing_attrs(style.find("w:pPr/w:spacing", NS))


def _style_spacing_by_name(docx_path, style_name):
    with zipfile.ZipFile(docx_path, "r") as zf:
        root = ET.fromstring(zf.read("word/styles.xml"))
    for style in root.findall("w:style", NS):
        name = style.find("w:name", NS)
        if name is not None and name.get(qn("w:val")) == style_name:
            return _spacing_attrs(style.find("w:pPr/w:spacing", NS))
    return {}


def _style_snap_to_grid(docx_path, style_id):
    with zipfile.ZipFile(docx_path, "r") as zf:
        root = ET.fromstring(zf.read("word/styles.xml"))
    style = root.find(f"w:style[@w:styleId='{style_id}']", NS)
    if style is None:
        return None
    snap = style.find("w:pPr/w:snapToGrid", NS)
    return None if snap is None else snap.get(qn("w:val"))


def _paragraph_spacing_by_text(docx_path, text):
    with zipfile.ZipFile(docx_path, "r") as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    for para in root.findall(".//w:p", NS):
        raw = "".join(node.text or "" for node in para.findall(".//w:t", NS))
        if raw == text:
            return _spacing_attrs(para.find("w:pPr/w:spacing", NS))
    raise AssertionError(f"Paragraph not found: {text}")


class LineSpacingModeTests(unittest.TestCase):
    def setUp(self):
        self.tmpdir = Path(tempfile.mkdtemp(prefix="line_spacing_test_", dir=PROJECT_ROOT / "tmp_test_artifacts"))

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def _run_apply(self, input_path, cfg=None, output_name=None):
        runtime_cfg = copy.deepcopy(cfg or DEFAULT_CONFIG)
        output_path = self.tmpdir / (output_name or f"{Path(input_path).stem}_output.docx")
        warnings = apply_format(str(input_path), str(output_path), config=runtime_cfg) or []
        return output_path, warnings, runtime_cfg

    def test_helper_normalization_supports_word_like_modes(self):
        self.assertEqual(normalize_line_spacing(1.5), {"mode": "multiple", "value": 1.5})
        self.assertEqual(normalize_line_spacing("1.5倍"), {"mode": "multiple", "value": 1.5})
        self.assertEqual(normalize_line_spacing("3倍"), {"mode": "multiple", "value": 3.0})
        self.assertEqual(normalize_line_spacing("固定值 20pt"), {"mode": "exact", "value": "20pt"})
        self.assertEqual(normalize_line_spacing("最小值 18pt"), {"mode": "at_least", "value": "18pt"})
        self.assertEqual(normalize_line_spacing("0.8cm"), {"mode": "exact", "value": "0.8cm"})
        self.assertEqual(format_line_spacing_value({"mode": "multiple", "value": 1.5}), "1.5倍")
        self.assertEqual(format_line_spacing_value({"mode": "multiple", "value": 3.0}), "3倍")
        self.assertEqual(format_line_spacing_value({"mode": "exact", "value": "20"}), "20pt")
        self.assertEqual(normalize_paragraph_spacing(0.5), {"mode": "lines", "value": 0.5})
        self.assertEqual(normalize_paragraph_spacing("12pt"), {"mode": "length", "value": "12pt"})
        self.assertEqual(format_paragraph_spacing_value(0.5), "0.5行")
        self.assertEqual(format_paragraph_spacing_value("12pt"), "12pt")
        self.assertEqual(line_spacing_to_ooxml(1.5), ("360", "auto"))
        self.assertEqual(line_spacing_to_ooxml(3.0), ("720", "auto"))
        self.assertEqual(line_spacing_to_ooxml({"mode": "exact", "value": "20pt"}), ("400", "exact"))
        self.assertEqual(line_spacing_to_ooxml({"mode": "at_least", "value": "18pt"}), ("360", "atLeast"))
        self.assertEqual(
            line_spacing_to_points_and_rule({"mode": "multiple", "value": 1.5}, base_line_pt=12),
            (18.0, int(WD_LINE_SPACING.MULTIPLE)),
        )
        self.assertEqual(
            line_spacing_to_points_and_rule({"mode": "multiple", "value": 3.0}, base_line_pt=12),
            (36.0, int(WD_LINE_SPACING.MULTIPLE)),
        )
        self.assertEqual(
            line_spacing_to_points_and_rule({"mode": "exact", "value": "20pt"}),
            (20.0, int(WD_LINE_SPACING.EXACTLY)),
        )
        self.assertEqual(
            line_spacing_to_points_and_rule({"mode": "at_least", "value": "18pt"}),
            (18.0, int(WD_LINE_SPACING.AT_LEAST)),
        )

    def test_gui_normalization_auto_adds_word_like_units(self):
        self.assertEqual(FormatterGUI._normalize_line_spacing_value("multiple", "1.5"), "1.5倍")
        self.assertEqual(FormatterGUI._normalize_line_spacing_value("multiple", "1.5倍"), "1.5倍")
        self.assertEqual(FormatterGUI._normalize_line_spacing_value("multiple", "3"), "3倍")
        self.assertEqual(FormatterGUI._normalize_line_spacing_value("multiple", "3倍"), "3倍")
        self.assertEqual(FormatterGUI._normalize_line_spacing_value("multiple", "20pt"), "1.5倍")
        self.assertEqual(FormatterGUI._normalize_line_spacing_value("exact", "20"), "20pt")
        self.assertEqual(FormatterGUI._normalize_line_spacing_value("exact", "1.5倍"), "20pt")
        self.assertEqual(FormatterGUI._normalize_line_spacing_value("exact", "0.8cm"), "0.8cm")
        self.assertEqual(FormatterGUI._normalize_line_spacing_value("at_least", "18"), "18pt")
        self.assertEqual(FormatterGUI._normalize_line_spacing_value("at_least", "0.8cm"), "0.8cm")
        self.assertEqual(FormatterGUI._collect_line_spacing_config("单倍行距", ""), 1.0)
        self.assertEqual(FormatterGUI._collect_line_spacing_config("1.5倍行距", ""), 1.5)
        self.assertEqual(FormatterGUI._collect_line_spacing_config("2倍行距", ""), 2.0)
        self.assertEqual(FormatterGUI._collect_line_spacing_config("多倍行距", "3"), 3.0)
        self.assertEqual(FormatterGUI._collect_line_spacing_config("多倍行距", "3倍"), 3.0)
        self.assertEqual(FormatterGUI._collect_line_spacing_config("多倍行距", "1.2"), 1.2)
        self.assertEqual(FormatterGUI._collect_line_spacing_config("多倍", "1.2"), 1.2)
        self.assertEqual(
            FormatterGUI._collect_line_spacing_config("固定值", "20"),
            {"mode": "exact", "value": "20pt"},
        )
        self.assertEqual(
            FormatterGUI._collect_line_spacing_config("最小值", "0.8cm"),
            {"mode": "at_least", "value": "0.8cm"},
        )
        self.assertEqual(FormatterGUI._split_line_spacing_for_gui(1.0), ("单倍行距", "1倍"))
        self.assertEqual(FormatterGUI._split_line_spacing_for_gui(1.5), ("1.5倍行距", "1.5倍"))
        self.assertEqual(FormatterGUI._split_line_spacing_for_gui(2.0), ("2倍行距", "2倍"))
        self.assertEqual(FormatterGUI._split_line_spacing_for_gui(3.0), ("多倍行距", "3倍"))
        self.assertEqual(FormatterGUI._split_line_spacing_for_gui(1.2), ("多倍行距", "1.2倍"))
        self.assertEqual(
            FormatterGUI._split_line_spacing_for_gui({"mode": "exact", "value": "20"}),
            ("固定值", "20pt"),
        )
        self.assertEqual(
            FormatterGUI._split_line_spacing_for_gui({"mode": "at_least", "value": "18pt"}),
            ("最小值", "18pt"),
        )

    def test_formatter_supports_exact_and_at_least_spacing_across_sections(self):
        doc = Document()
        doc.add_paragraph("第1章 绪论", style="Heading 1")
        doc.add_paragraph("正文行距测试")
        doc.add_picture(str(IMAGE_PATH))
        doc.add_paragraph("图1.1 题注行距测试")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "表格行距测试"
        doc.styles.add_style("Footnote Text", WD_STYLE_TYPE.PARAGRAPH)
        input_path = self.tmpdir / "mixed_spacing_input.docx"
        doc.save(input_path)

        cfg = copy.deepcopy(DEFAULT_CONFIG)
        cfg["cover"]["enabled"] = False
        cfg["body"]["line_spacing"] = {"mode": "exact", "value": "20pt"}
        cfg["captions"]["line_spacing"] = {"mode": "at_least", "value": "18pt"}
        cfg["table"]["line_spacing"] = {"mode": "exact", "value": "14pt"}
        cfg["footnote"]["line_spacing"] = {"mode": "at_least", "value": "16pt"}

        output_path, warnings, _ = self._run_apply(input_path, cfg=cfg)

        self.assertTrue(output_path.exists())
        self.assertIsInstance(warnings, list)
        self.assertEqual(
            _paragraph_spacing_by_text(output_path, "正文行距测试"),
            {"line": "400", "lineRule": "exact", "before": None, "beforeLines": None, "after": None, "afterLines": None},
        )
        self.assertEqual(
            _paragraph_spacing_by_text(output_path, "图1.1 题注行距测试"),
            {"line": "360", "lineRule": "atLeast", "before": None, "beforeLines": None, "after": "240", "afterLines": None},
        )
        self.assertEqual(
            _paragraph_spacing_by_text(output_path, "表格行距测试"),
            {"line": "280", "lineRule": "exact", "before": None, "beforeLines": None, "after": None, "afterLines": None},
        )
        self.assertEqual(
            _style_spacing_by_name(output_path, "Footnote Text"),
            {"line": "320", "lineRule": "atLeast", "before": None, "beforeLines": None, "after": None, "afterLines": None},
        )

    def test_toc_styles_support_multiple_exact_and_at_least_spacing(self):
        doc = Document()
        doc.add_paragraph("第1章 绪论", style="Heading 1")
        doc.add_paragraph("1.1 研究背景", style="Heading 2")
        input_path = self.tmpdir / "toc_spacing_modes_input.docx"
        doc.save(input_path)

        cases = [
            ("multiple", 2.0, {"line": "480", "lineRule": "auto"}),
            ("multiple_3x", 3.0, {"line": "720", "lineRule": "auto"}),
            ("exact", {"mode": "exact", "value": "20pt"}, {"line": "400", "lineRule": "exact"}),
            ("at_least", {"mode": "at_least", "value": "18pt"}, {"line": "360", "lineRule": "atLeast"}),
        ]

        for label, line_spacing, expected_line in cases:
            with self.subTest(mode=label):
                cfg = copy.deepcopy(DEFAULT_CONFIG)
                cfg["cover"]["enabled"] = False
                cfg["toc"]["line_spacing"] = line_spacing
                cfg["toc"]["space_before"] = 1.0
                cfg["toc"]["space_after"] = 0.5

                output_path, _, _ = self._run_apply(
                    input_path,
                    cfg=cfg,
                    output_name=f"toc_{label}.docx",
                )

                expected_toc1 = {
                    **expected_line,
                    "before": None,
                    "beforeLines": str(int(cfg["toc"]["space_before"] * 100)),
                    "after": None,
                    "afterLines": str(int(cfg["toc"]["space_after"] * 100)),
                }
                expected_toc2 = {
                    **expected_line,
                    "before": None,
                    "beforeLines": str(int(cfg["toc"]["space_before"] * 100)),
                    "after": None,
                    "afterLines": str(int(cfg["toc"]["space_after"] * 100)),
                }
                self.assertEqual(_style_spacing_by_id(output_path, "TOC1"), expected_toc1)
                self.assertEqual(_style_spacing_by_id(output_path, "TOC2"), expected_toc2)
                self.assertEqual(_style_snap_to_grid(output_path, "TOC1"), "0")
                self.assertEqual(_style_snap_to_grid(output_path, "TOC2"), "0")


if __name__ == "__main__":
    unittest.main()
