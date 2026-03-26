import tempfile
import unittest
from pathlib import Path

import yaml
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

from thesis_config import dump_default_config
from thesis_formatter.formatter import apply_format


def _make_config():
    cfg = yaml.safe_load(dump_default_config())
    cfg["cover"]["enabled"] = False
    cfg["toc"]["enabled"] = False
    cfg["header_footer"]["enabled"] = True
    cfg["header_footer"]["scope"] = "body"
    cfg["header_footer"]["odd_page_text"] = "ODD_HEADER"
    cfg["header_footer"]["even_page_text"] = "EVEN_HEADER"
    cfg["header_footer"]["different_odd_even"] = True
    cfg["header_footer"]["border_bottom"] = False
    cfg["page_numbers"]["front_position"] = "center"
    cfg["page_numbers"]["body_position"] = "alternate"
    cfg["page_numbers"]["body_odd_position"] = "right"
    cfg["page_numbers"]["body_even_position"] = "left"
    return cfg


def _heading_1(doc, text):
    para = doc.add_paragraph(text)
    para.style = doc.styles["Heading 1"]
    return para


def _build_doc_with_front_section_and_body_heading_in_next_section():
    doc = Document()
    doc.add_paragraph("摘要")
    doc.add_paragraph("这是中文摘要内容。")
    doc.add_paragraph("关键词：测试；分页；页码")

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("目录")
    doc.add_paragraph("第一章 绪论 ................................ 1")
    doc.add_paragraph("第二章 方法 ................................ 5")

    _heading_1(doc, "第1章 绪论")
    doc.add_paragraph("这是正文第一页。")
    return doc


def _section_pg_num_attrs(section):
    node = section._sectPr.find(qn("w:pgNumType"))
    if node is None:
        return {}
    return {key.split("}", 1)[-1]: value for key, value in node.attrib.items()}


class StandaloneHeaderPageNumberTests(unittest.TestCase):
    def test_page_number_only_mode_keeps_existing_section_count(self):
        cfg = _make_config()
        cfg["page_numbers"]["only_insert"] = True
        cfg["header_footer"]["enabled"] = False

        with tempfile.TemporaryDirectory() as tmp:
            input_path = Path(tmp) / "input.docx"
            output_path = Path(tmp) / "output.docx"

            _build_doc_with_front_section_and_body_heading_in_next_section().save(input_path)
            apply_format(str(input_path), str(output_path), config=cfg)

            out = Document(output_path)

        self.assertEqual(len(out.sections), 2)
        self.assertEqual(
            _section_pg_num_attrs(out.sections[0]),
            {"fmt": "upperRoman", "start": "1"},
        )
        self.assertEqual(
            _section_pg_num_attrs(out.sections[1]),
            {"fmt": "decimal", "start": "1"},
        )

    def test_header_only_mode_uses_existing_sections_without_inserting_new_ones(self):
        cfg = _make_config()
        cfg["header_footer"]["only_insert"] = True

        with tempfile.TemporaryDirectory() as tmp:
            input_path = Path(tmp) / "input.docx"
            output_path = Path(tmp) / "output.docx"

            _build_doc_with_front_section_and_body_heading_in_next_section().save(input_path)
            apply_format(str(input_path), str(output_path), config=cfg)

            out = Document(output_path)

        self.assertEqual(len(out.sections), 2)
        self.assertEqual(out.sections[0].header.paragraphs[0].text, "")
        self.assertEqual(out.sections[0].even_page_header.paragraphs[0].text, "")
        self.assertEqual(out.sections[1].header.paragraphs[0].text, "ODD_HEADER")
        self.assertEqual(out.sections[1].even_page_header.paragraphs[0].text, "EVEN_HEADER")

    def test_multiple_local_modes_raise_runtime_error(self):
        cfg = _make_config()
        cfg["toc"]["only_insert"] = True
        cfg["page_numbers"]["only_insert"] = True

        with tempfile.TemporaryDirectory() as tmp:
            input_path = Path(tmp) / "input.docx"
            output_path = Path(tmp) / "output.docx"

            _build_doc_with_front_section_and_body_heading_in_next_section().save(input_path)
            with self.assertRaisesRegex(RuntimeError, "单独处理模式不能同时启用多个"):
                apply_format(str(input_path), str(output_path), config=cfg)


if __name__ == "__main__":
    unittest.main()
