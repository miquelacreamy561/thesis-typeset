import copy
import os
import re
import sys
import tempfile
import zipfile
import subprocess

from thesis_config import resolve_config
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from ._common import (
    _ALIGN_MAP, set_style_font, set_run_font, set_para_runs_font,
    set_table_border, clear_table_border,
    _ensure_keep_next, _set_para_spacing, _check_caption_numbering,
    is_heading, contains_cjk, normalize_cn_keywords, normalize_en_keywords,
    get_paragraph_heading_level, get_heading_style, is_heading_style,
    _ALL_HEADING_NAMES, parse_length,
)
from ._titles import _find_special_display, _get_special_title_map, _detect_front_matter
from .headings import auto_assign_heading_styles, renumber_headings, normalize_heading_spacing
from .references import check_citations, apply_ref_crosslinks
from .page import normalize_sections, setup_page_numbers, insert_page_break_after, find_first_body_heading
from .headers import setup_headers
from .toc import insert_toc, ensure_toc_styles
from .cover import _has_cover, find_existing_cover_end, insert_cover_and_declaration
from .structure import validate_structure


def _insert_cover_via_vbs(target_path, cover_path):
    """使用嵌入的 VBS 代码插入封面，保留完整格式."""
    import sys

    vbs_code = """Option Explicit
Dim objWord, targetDoc, args, targetPath, coverPath
Set args = WScript.Arguments
If args.Count < 2 Then WScript.Quit 1
targetPath = args(0): coverPath = args(1)
On Error Resume Next
Set objWord = CreateObject("Word.Application")
If Err.Number <> 0 Then WScript.Quit 1
On Error GoTo 0
objWord.Visible = False: objWord.DisplayAlerts = 0
Set targetDoc = objWord.Documents.Open(targetPath)
If Err.Number <> 0 Then objWord.Quit: WScript.Quit 1
On Error GoTo 0
objWord.Selection.HomeKey 6
objWord.Selection.InsertFile coverPath
objWord.Selection.EndKey 6
objWord.Selection.InsertBreak 7
targetDoc.Save
targetDoc.Close False
objWord.Quit
WScript.Echo "Done"
"""

    try:
        import tempfile
        with tempfile.NamedTemporaryFile(mode="w", suffix=".vbs", delete=False) as f:
            f.write(vbs_code)
            vbs_path = f.name

        result = subprocess.run(
            ["cscript", "//NoLogo", vbs_path, target_path, cover_path],
            capture_output=True,
            text=True,
            encoding="gbk",
            timeout=60
        )

        os.unlink(vbs_path)
        return result.returncode == 0, result.stderr if result.returncode != 0 else ""
    except Exception as e:
        return False, str(e)


def apply_format(input_path, output_path, config=None, config_path=None):
    if config is None:
        config, config_path = resolve_config(input_path=input_path)
    cfg = config

    latin = cfg["fonts"]["latin"]
    body_font = cfg["fonts"]["body"]
    body_size = parse_length(cfg["sizes"]["body"])
    body_ls = cfg["body"]["line_spacing"]
    body_indent = parse_length(cfg["body"]["first_line_indent"])
    body_align = _ALIGN_MAP.get(cfg["body"]["align"])

    h1_font = cfg["fonts"]["h1"]
    h1_size = parse_length(cfg["sizes"]["h1"])
    h2_font = cfg["fonts"]["h2"]
    h2_size = parse_length(cfg["sizes"]["h2"])
    h3_font = cfg["fonts"]["h3"]
    h3_size = parse_length(cfg["sizes"]["h3"])
    h4_font = cfg["fonts"]["h4"]
    h4_size = parse_length(cfg["sizes"]["h4"])

    def _bold_val(v):
        return None if v == "keep" else v

    h1_bold = _bold_val(cfg["headings"]["h1"]["bold"])
    h1_align = _ALIGN_MAP.get(cfg["headings"]["h1"]["align"])
    h2_bold = _bold_val(cfg["headings"]["h2"]["bold"])
    h2_align = _ALIGN_MAP.get(cfg["headings"]["h2"]["align"])
    h3_bold = _bold_val(cfg["headings"]["h3"]["bold"])
    h3_align = _ALIGN_MAP.get(cfg["headings"]["h3"]["align"])
    h4_bold = _bold_val(cfg["headings"]["h4"]["bold"])
    h4_align = _ALIGN_MAP.get(cfg["headings"]["h4"]["align"])

    caption_size = parse_length(cfg["sizes"]["caption"])
    note_size = parse_length(cfg["sizes"]["note"])
    fn_size = parse_length(cfg["sizes"]["footnote"])

    st_map = _get_special_title_map(cfg)
    sec = cfg["sections"]
    ref_key = "\u53c2\u8003\u6587\u732e"
    toc_key = "\u76ee\u5f55"

    doc = Document(input_path)

    fm_mode = cfg.get("front_matter", {}).get("mode", "auto")

    preserve_until_idx = 0
    first_body_heading = find_first_body_heading(doc, cfg)
    if fm_mode == "skip":
        if first_body_heading is None:
            preserve_until_idx = len(doc.paragraphs)
        else:
            for idx, para in enumerate(doc.paragraphs):
                if para._element is first_body_heading._element:
                    preserve_until_idx = idx
                    break
    elif not cfg.get("cover", {}).get("enabled", True):
        preserve_until_idx = find_existing_cover_end(doc, cfg)

    preserved_para_ids = {id(para._element) for para in doc.paragraphs[:preserve_until_idx]}
    preserve_front_matter = bool(preserved_para_ids)
    preserved_front_snapshots = [
        copy.deepcopy(para._element)
        for para in doc.paragraphs[:preserve_until_idx]
    ] if preserve_front_matter else []

    def _is_preserved_front_para(para):
        return id(para._element) in preserved_para_ids

    def _restore_preserved_front_paragraphs():
        if not preserved_front_snapshots:
            return
        for idx, original_el in enumerate(preserved_front_snapshots):
            if idx >= len(doc.paragraphs):
                break

            current_el = doc.paragraphs[idx]._element
            parent = current_el.getparent()
            if parent is None:
                continue

            current_ppr = current_el.find(qn("w:pPr"))
            current_sectpr = None
            if current_ppr is not None:
                sectpr = current_ppr.find(qn("w:sectPr"))
                if sectpr is not None:
                    current_sectpr = copy.deepcopy(sectpr)

            restored_el = copy.deepcopy(original_el)
            if current_sectpr is not None:
                restored_ppr = restored_el.find(qn("w:pPr"))
                if restored_ppr is None:
                    restored_ppr = OxmlElement("w:pPr")
                    restored_el.insert(0, restored_ppr)
                for old_sectpr in list(restored_ppr.findall(qn("w:sectPr"))):
                    restored_ppr.remove(old_sectpr)
                restored_ppr.append(current_sectpr)

            parent.replace(current_el, restored_el)

    cover_end_idx = preserve_until_idx

    warnings = []
    from . import numbering
    requested_caption_mode, raw_caption_mode, raw_caption_mode_warnings, raw_caption_reasons = numbering.resolve_caption_mode(doc, cfg)
    cfg.setdefault("_runtime", {})
    cfg["_runtime"]["caption_mode_requested"] = requested_caption_mode
    cfg["_runtime"]["caption_mode_raw_effective"] = raw_caption_mode
    cfg["_runtime"]["caption_mode_raw_precheck_passed"] = requested_caption_mode != numbering.CAPTION_MODE_DYNAMIC or not raw_caption_reasons

    auto_changes = auto_assign_heading_styles(doc, cfg, skip_para_ids=preserved_para_ids)
    if auto_changes:
        print(f"\u81ea\u52a8\u8bc6\u522b\u6807\u9898 ({len(auto_changes)} \u4e2a):", file=sys.stderr)
        for c in auto_changes:
            print(c, file=sys.stderr)

    try:
        warnings.extend(validate_structure(doc, cfg) or [])
    except Exception as exc:
        print(f"\u7ed3\u6784\u68c0\u67e5\u51fa\u9519\uff08\u5df2\u8df3\u8fc7\uff0c\u7ee7\u7eed\u6392\u7248\uff09: {exc}", file=sys.stderr)
    normalize_sections(doc, cfg)

    renum_changes = []
    if sec.get("renumber_headings", False):
        renum_changes = numbering.setup_multilevel_list(doc, cfg)

    requested_caption_mode, effective_caption_mode, caption_mode_warnings = numbering.resolve_caption_mode_after_normalization(
        doc, cfg, raw_reasons=raw_caption_reasons
    )
    cfg["_runtime"]["caption_mode_effective"] = effective_caption_mode
    cfg["_runtime"]["caption_mode_precheck_passed"] = effective_caption_mode == numbering.CAPTION_MODE_DYNAMIC
    warnings.extend(raw_caption_mode_warnings)
    warnings.extend(caption_mode_warnings)

    normalize_heading_spacing(doc, cfg, skip_para_ids=preserved_para_ids)

    if not preserve_front_matter:
        for style_name in ["Normal", "Body Text", "First Paragraph", "_Style 2"]:
            if style_name in doc.styles:
                set_style_font(doc.styles[style_name], east_asia=body_font,
                               size_pt=body_size, bold=False, latin=latin)

    def _set_heading_style(level, font, size, bold, align, hcfg):
        style = get_heading_style(doc, level)
        if style is None:
            return
        set_style_font(style, east_asia=font, size_pt=size, bold=bold, latin=latin)
        if align is not None:
            style.paragraph_format.alignment = align
        sb = hcfg.get("space_before", 0)
        sa = hcfg.get("space_after", 0)
        if sb >= 0:
            style.paragraph_format.space_before = parse_length(sb * 12)
        if sa >= 0:
            style.paragraph_format.space_after = parse_length(sa * 12)

    if not preserve_front_matter:
        _set_heading_style(1, h1_font, h1_size, h1_bold, h1_align, cfg["headings"]["h1"])
        _set_heading_style(2, h2_font, h2_size, h2_bold, h2_align, cfg["headings"]["h2"])
        _set_heading_style(3, h3_font, h3_size, h3_bold, h3_align, cfg["headings"]["h3"])
        _set_heading_style(4, h4_font, h4_size, h4_bold, h4_align, cfg["headings"]["h4"])

    if "TOC Heading" in doc.styles:
        st = doc.styles["TOC Heading"]
        toc_h_font = cfg["toc"].get("h1_font", h1_font)
        toc_h_size = parse_length(cfg["toc"].get("h1_font_size", cfg["sizes"]["h1"]))
        set_style_font(st, east_asia=toc_h_font, size_pt=toc_h_size, bold=True, latin=latin)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _toc_h_sb = cfg["toc"].get("space_before", 0)
        _toc_h_sa = cfg["toc"].get("space_after", 0)
        st.paragraph_format.space_before = parse_length(_toc_h_sb * 12)
        st.paragraph_format.space_after = parse_length(_toc_h_sa * 12)

    ensure_toc_styles(doc, cfg)

    toc_content_font = cfg["toc"].get("font", body_font)
    toc_content_size = parse_length(cfg["toc"].get("font_size", cfg["sizes"]["body"]))
    toc_h1_font = cfg["toc"].get("h1_font", cfg["fonts"]["h1"])
    toc_h1_size = parse_length(cfg["toc"].get("h1_font_size", cfg["sizes"]["h1"]))
    toc_content_ls = cfg["toc"].get("line_spacing", body_ls)
    toc_sb = parse_length(cfg["toc"].get("space_before", 0) * 12)
    toc_sa = parse_length(cfg["toc"].get("space_after", 0) * 12)

    for para in doc.paragraphs:
        if _is_preserved_front_para(para):
            continue
        sn = para.style.name if para.style else ""
        if sn.lower().startswith("toc ") or sn == "\u6837\u5f0f3":
            is_toc1 = sn.lower() == "toc 1"
            para.paragraph_format.first_line_indent = parse_length(0)
            para.paragraph_format.line_spacing = toc_content_ls
            para.paragraph_format.space_before = toc_sb
            para.paragraph_format.space_after = toc_sa
            ea = toc_h1_font if is_toc1 else toc_content_font
            sz = toc_h1_size if is_toc1 else toc_content_size
            set_para_runs_font(para, east_asia=ea, size_pt=sz,
                               bold=False, latin=latin)

    for name in ["Footnote Text", "Footnote Reference"]:
        if name in doc.styles:
            set_style_font(doc.styles[name], east_asia=body_font, size_pt=fn_size,
                           bold=False, latin=latin)
    if "Footnote Text" in doc.styles:
        ft = doc.styles["Footnote Text"]
        ft.paragraph_format.line_spacing = cfg["footnote"]["line_spacing"]
        ft.paragraph_format.first_line_indent = parse_length(0)
        _fn_align = _ALIGN_MAP.get(cfg["footnote"].get("align", "justify"))
        if _fn_align is not None:
            ft.paragraph_format.alignment = _fn_align

    for name in ["Hyperlink", "\u8d85\u94fe\u63a5"]:
        if name in [s.name for s in doc.styles]:
            st = doc.styles[name]
            st.font.color.rgb = RGBColor(0, 0, 0)
            st.font.underline = False

    for para in doc.paragraphs:
        if _is_preserved_front_para(para):
            continue
        level = get_paragraph_heading_level(para)
        sn = para.style.name if para.style else ""
        if sn.lower().startswith("toc ") or sn == "\u6837\u5f0f3":
            continue
        if level is not None:
            hkey = {1: "h1", 2: "h2", 3: "h3", 4: "h4"}.get(level, "h1")
            hcfg = cfg["headings"].get(hkey, {})
            if hcfg.get("space_before", 0) >= 0:
                para.paragraph_format.space_before = parse_length(0)
            if hcfg.get("space_after", 0) >= 0:
                para.paragraph_format.space_after = parse_length(0)
        else:
            para.paragraph_format.space_before = parse_length(cfg["body"].get("space_before", 0) * 12)
            para.paragraph_format.space_after = parse_length(cfg["body"].get("space_after", 0) * 12)
        pf = para.paragraph_format
        if para.style and para.style.name in ["Normal", "Body Text", "First Paragraph", "_Style 2"]:
            if body_align is not None:
                pf.alignment = body_align
            pf.first_line_indent = body_indent
            pf.line_spacing = body_ls
            set_para_runs_font(para, east_asia=body_font, size_pt=body_size,
                               bold=False, latin=latin)

    fm_mode = cfg.get("front_matter", {}).get("mode", "auto")
    has_fm = (fm_mode == "format") or \
             (fm_mode == "auto" and _detect_front_matter(doc, cfg))

    cn_kw_para = None
    en_kw_para = None

    if has_fm:
        first_h1_idx = None
        for i, para in enumerate(doc.paragraphs):
            if is_heading(para, 1):
                first_h1_idx = i
                break
        if first_h1_idx is None:
            first_h1_idx = len(doc.paragraphs)

        front = doc.paragraphs[cover_end_idx:first_h1_idx]
        non_empty = [p for p in front if p.text.strip()]

        if non_empty:
            abstract_display = _find_special_display(cfg, "\u6458\u8981")
            p = non_empty[0]
            p.text = abstract_display
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = parse_length(0)
            p.paragraph_format.line_spacing = body_ls
            set_para_runs_font(p, east_asia=h1_font, size_pt=h1_size, bold=True, latin=latin)

        cn_kw_re = sec.get("cn_keywords_pattern", r"^\s*\u5173\u952e\u8bcd\s*[\uff1a:]")
        en_abs_re = sec.get("en_abstract_pattern", r"(?i)^\s*Abstract\s*[\uff1a:]")
        en_kw_re = sec.get("en_keywords_pattern", r"(?i)^\s*Key\s*words\s*[\uff1a:]")

        past_abstract = False
        en_title_seen = False

        for p in non_empty[1:]:
            t = p.text.strip()
            if t.startswith("\u5173\u952e\u8bcd"):
                cn_kw_para = p
                normalized = normalize_cn_keywords(t) or t
                content = normalized.split("\uff1a", 1)[1] if "\uff1a" in normalized else ""
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.first_line_indent = parse_length(0)
                p.paragraph_format.line_spacing = body_ls
                r1 = p.add_run("\u5173\u952e\u8bcd\uff1a")
                set_run_font(r1, east_asia=h1_font, size_pt=body_size, bold=True, latin=latin)
                r2 = p.add_run(content)
                set_run_font(r2, east_asia=body_font, size_pt=body_size, bold=False, latin=latin)
            elif re.match(r"^\s*Abstract\s*:", t, flags=re.I):
                past_abstract = True
                content = re.sub(r"^\s*Abstract\s*:\s*", "", t, flags=re.I)
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = parse_length(0)
                p.paragraph_format.line_spacing = body_ls
                r1 = p.add_run("Abstract: ")
                set_run_font(r1, east_asia=latin, size_pt=body_size, bold=True, latin=latin)
                r2 = p.add_run(content)
                set_run_font(r2, east_asia=latin, size_pt=body_size, bold=False, latin=latin)
            elif re.match(r"^\s*Key\s*words\s*:", t, flags=re.I):
                en_kw_para = p
                normalized = normalize_en_keywords(t) or t
                content = re.sub(r"^\s*Key\s*words\s*:\s*", "", normalized, flags=re.I)
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = parse_length(0)
                p.paragraph_format.line_spacing = body_ls
                r1 = p.add_run("Key words: ")
                set_run_font(r1, east_asia=latin, size_pt=body_size, bold=True, latin=latin)
                r2 = p.add_run(content)
                set_run_font(r2, east_asia=latin, size_pt=body_size, bold=False, latin=latin)
            elif not past_abstract and not contains_cjk(t) and not re.match(r"^\s*(Abstract|Key\s*words)\s*:", t, re.I) and len(t) > 20 and not re.match(r"^[\(\uff08]", t):
                en_title_seen = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = parse_length(0)
                p.paragraph_format.line_spacing = body_ls
                set_para_runs_font(p, east_asia=latin, size_pt=h1_size, bold=True, latin=latin)
            elif not past_abstract and en_title_seen and not contains_cjk(t) and not re.match(r"^[\(\uff08]", t) and not re.match(r"^\s*(Abstract|Key\s*words)\s*:", t, re.I):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = parse_length(0)
                p.paragraph_format.line_spacing = body_ls
                set_para_runs_font(p, east_asia=latin, size_pt=body_size, bold=False, latin=latin)
            elif not past_abstract and re.match(r"^[\(\uff08]", t) and re.search(r"(China|University|College)", t, re.I):
                new_t = t
                if new_t.startswith("("):
                    new_t = "\uff08" + new_t[1:]
                if new_t.endswith(")"):
                    new_t = new_t[:-1] + "\uff09"
                if new_t != t:
                    p.text = new_t
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = parse_length(0)
                p.paragraph_format.line_spacing = body_ls
                set_para_runs_font(p, east_asia=latin, size_pt=body_size, bold=False, latin=latin)
            else:
                if contains_cjk(t):
                    if body_align is not None:
                        p.alignment = body_align
                    p.paragraph_format.first_line_indent = body_indent
                    p.paragraph_format.line_spacing = body_ls
                    set_para_runs_font(p, east_asia=body_font, size_pt=body_size,
                                       bold=False, latin=latin)
                elif t:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.first_line_indent = parse_length(0)
                    p.paragraph_format.line_spacing = body_ls
                    set_para_runs_font(p, east_asia=latin, size_pt=body_size,
                                       bold=False, latin=latin)

        if cn_kw_para is not None:
            insert_page_break_after(cn_kw_para)
        if en_kw_para is not None:
            insert_page_break_after(en_kw_para)

    def _apply_heading_para(para, align, hcfg, font, size, bold):
        if align is not None:
            para.alignment = align
        para.paragraph_format.first_line_indent = parse_length(0)
        para.paragraph_format.line_spacing = body_ls
        sb = hcfg.get("space_before", 0)
        sa = hcfg.get("space_after", 0)
        if sb >= 0:
            para.paragraph_format.space_before = parse_length(sb * 12)
        if sa >= 0:
            para.paragraph_format.space_after = parse_length(sa * 12)
        set_para_runs_font(para, east_asia=font, size_pt=size, bold=bold, latin=latin)

    for para in doc.paragraphs:
        if _is_preserved_front_para(para):
            continue
        level = get_paragraph_heading_level(para)
        t = para.text.strip()
        t_nospace = t.replace(" ", "").replace("\u3000", "")
        sn = para.style.name if para.style else ""

        if level is not None and t_nospace in st_map:
            entry = st_map[t_nospace]
            para.text = entry["display"]
            para.alignment = _ALIGN_MAP.get(entry.get("align", "center"), WD_ALIGN_PARAGRAPH.CENTER)
            para.paragraph_format.first_line_indent = parse_length(0)
            para.paragraph_format.line_spacing = body_ls
            _h1cfg = cfg["headings"]["h1"]
            _sb = _h1cfg.get("space_before", 0)
            _sa = _h1cfg.get("space_after", 0)
            if _sb >= 0:
                para.paragraph_format.space_before = parse_length(_sb * 12)
            if _sa >= 0:
                para.paragraph_format.space_after = parse_length(_sa * 12)
            set_para_runs_font(para, east_asia=h1_font, size_pt=h1_size,
                               bold=True, latin=latin)
        elif level == 1:
            _apply_heading_para(para, h1_align, cfg["headings"]["h1"], h1_font, h1_size, h1_bold)
            if t_nospace in st_map:
                entry = st_map[t_nospace]
                para.text = entry["display"]
                para.alignment = _ALIGN_MAP.get(entry.get("align", "center"),
                                                WD_ALIGN_PARAGRAPH.CENTER)
                set_para_runs_font(para, east_asia=h1_font, size_pt=h1_size,
                                   bold=True, latin=latin)
            elif t.startswith("\u9644\u5f55"):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_para_runs_font(para, east_asia=h1_font, size_pt=h1_size,
                                   bold=True, latin=latin)
        elif level == 2:
            _apply_heading_para(para, h2_align, cfg["headings"]["h2"], h2_font, h2_size, h2_bold)
        elif level == 3:
            _apply_heading_para(para, h3_align, cfg["headings"]["h3"], h3_font, h3_size, h3_bold)
        elif level == 4:
            _apply_heading_para(para, h4_align, cfg["headings"]["h4"], h4_font, h4_size, h4_bold)
        elif level == 1:
            para.paragraph_format.first_line_indent = parse_length(0)
            para.paragraph_format.line_spacing = body_ls
            sb = cfg["headings"]["h1"].get("space_before", 0)
            sa = cfg["headings"]["h1"].get("space_after", 0)
            if sb >= 0:
                para.paragraph_format.space_before = parse_length(sb * 12)
            if sa >= 0:
                para.paragraph_format.space_after = parse_length(sa * 12)
            if t_nospace in st_map:
                entry = st_map[t_nospace]
                para.text = entry["display"]
                para.alignment = _ALIGN_MAP.get(entry.get("align", "center"),
                                                WD_ALIGN_PARAGRAPH.CENTER)
                set_para_runs_font(para, east_asia=h1_font, size_pt=h1_size,
                                   bold=True, latin=latin)
            elif t.startswith("\u9644\u5f55"):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_para_runs_font(para, east_asia=h1_font, size_pt=h1_size,
                                   bold=True, latin=latin)

    ref_cfg = cfg["references"]
    in_refs = False
    for para in doc.paragraphs:
        if _is_preserved_front_para(para):
            continue
        level = get_paragraph_heading_level(para)
        t = para.text.strip().replace(" ", "").replace("\u3000", "")
        is_h1_like = level == 1
        if is_h1_like and t == ref_key.replace(" ", "").replace("\u3000", ""):
            in_refs = True
            continue
        if is_h1_like and in_refs:
            in_refs = False
        if in_refs and para.text.strip() and not (para.style and is_heading_style(para.style)):
            para.paragraph_format.first_line_indent = parse_length(ref_cfg["first_line_indent"])
            para.paragraph_format.left_indent = parse_length(ref_cfg["left_indent"])
            para.paragraph_format.line_spacing = body_ls
            set_para_runs_font(para, east_asia=body_font, size_pt=body_size,
                               bold=False, latin=latin)

    cap_cfg = cfg.get("captions", {})
    fig_pat = cap_cfg.get("figure_pattern", r"^\u56fe\s*\d")
    tbl_pat = cap_cfg.get("table_pattern", r"^(\u7eed)?\u8868\s*\d")
    subfig_pat = cap_cfg.get("subfigure_pattern", r"^\([a-z]\)")
    note_pat = cap_cfg.get("note_pattern", r"^\u6ce8[\uff1a:]")

    cap_ls = cap_cfg.get("line_spacing", body_ls)
    # 跳过已使用 SEQ 域的段落，只处理纯文本的题注
    _cap_space_re = re.compile(r"^((?:\u56fe|\u8868|Figure|Table)\s*[A-Z]?\d+)(\S)", re.I)
    for para in doc.paragraphs:
        if _is_preserved_front_para(para):
            continue
        # 跳过包含 SEQ 域的段落
        has_seq = False
        for run in para.runs:
            if "instrText" in str(run._element):
                has_seq = True
                break
        if has_seq:
            continue

        t = para.text.strip()
        m = _cap_space_re.match(t)
        if m:
            para.text = m.group(1) + " " + t[m.end(1):]

    spacing_line = parse_length(cfg["sizes"]["body"])
    source_pat = r"^(\u8d44\u6599)?\u6765\u6e90\s*[\uff1a:]"

    for para in doc.paragraphs:
        if _is_preserved_front_para(para):
            continue
        t = para.text.strip()
        if re.match(fig_pat, t) or re.match(r"^Figure\s*\d", t, re.I) or re.match(r"^\u56fe[A-Z]\d+", t):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = parse_length(0)
            para.paragraph_format.line_spacing = cap_ls
            para.paragraph_format.space_after = spacing_line
            set_para_runs_font(para, east_asia=body_font, size_pt=caption_size,
                               bold=False, latin=latin)
        elif re.match(tbl_pat, t) or re.match(r"^Table\s*\d", t, re.I) or re.match(r"^(\u7eed)?\u8868[A-Z]\d+", t):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = parse_length(0)
            para.paragraph_format.line_spacing = cap_ls
            para.paragraph_format.space_before = spacing_line
            set_para_runs_font(para, east_asia=body_font, size_pt=caption_size,
                               bold=False, latin=latin)
        elif re.match(subfig_pat, t):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = parse_length(0)
            para.paragraph_format.line_spacing = cap_ls
            set_para_runs_font(para, east_asia=body_font, size_pt=caption_size,
                               bold=False, latin=latin)
        elif re.match(note_pat, t) or re.match(source_pat, t):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = parse_length(0)
            para.paragraph_format.line_spacing = cap_ls
            set_para_runs_font(para, east_asia=body_font, size_pt=note_size,
                               bold=False, latin=latin)

    if cap_cfg.get("keep_with_next", True):
        body_el = doc.element.body
        children = list(body_el)
        for i, el in enumerate(children):
            if el.tag == qn("w:tbl"):
                if i > 0 and children[i - 1].tag == qn("w:p"):
                    prev_text = "".join(
                        (nd.text or "") for nd in children[i - 1].iter(qn("w:t"))
                    ).strip()
                    if re.match(tbl_pat, prev_text) or re.match(r"^Table\s*\d", prev_text, re.I) or re.match(r"^(\u7eed)?\u8868[A-Z]\d+", prev_text):
                        _ensure_keep_next(children[i - 1])
                if i + 1 < len(children) and children[i + 1].tag == qn("w:p"):
                    nt = "".join(
                        (nd.text or "") for nd in children[i + 1].iter(qn("w:t"))
                    ).strip()
                    if nt and not re.match(note_pat, nt) and not re.match(source_pat, nt):
                        _set_para_spacing(children[i + 1], "before", spacing_line)
            elif el.tag == qn("w:p") and el.findall(".//" + qn("w:drawing")):
                if i + 1 < len(children) and children[i + 1].tag == qn("w:p"):
                    next_text = "".join(
                        (nd.text or "") for nd in children[i + 1].iter(qn("w:t"))
                    ).strip()
                    if re.match(fig_pat, next_text) or re.match(r"^Figure\s*\d", next_text, re.I) or re.match(r"^\u56fe[A-Z]\d+", next_text):
                        _ensure_keep_next(el)
                if i > 0 and children[i - 1].tag == qn("w:p"):
                    pt = "".join(
                        (nd.text or "") for nd in children[i - 1].iter(qn("w:t"))
                    ).strip()
                    if pt and not re.match(fig_pat, pt) and not re.match(subfig_pat, pt):
                        _set_para_spacing(el, "before", spacing_line)

    if cap_cfg.get("check_numbering", True):
        warnings.extend(_check_caption_numbering(doc, fig_pat, tbl_pat, cfg))

    if cap_cfg.get("use_seq_fields", True):
        from . import numbering
        numbering.setup_figure_captions(doc, cfg)
        numbering.setup_table_captions(doc, cfg)

    try:
        warnings.extend(check_citations(doc, cfg))
    except Exception as exc:
        print(f"\u5f15\u7528\u68c0\u67e5\u51fa\u9519\uff08\u5df2\u8df3\u8fc7\uff09: {exc}", file=sys.stderr)

    _cite_comma = re.compile(r",\s*((?:19|20)\d{2})")
    for para in doc.paragraphs:
        if _is_preserved_front_para(para):
            continue
        for run in para.runs:
            old = run.text
            new = _cite_comma.sub(r", \1", old)
            if new != old:
                run.text = new
                print(f"  \u5f15\u7528\u9017\u53f7\u4fee\u6b63: \"{old.strip()[:40]}\" \u2192 \"{new.strip()[:40]}\"")

    try:
        apply_ref_crosslinks(doc, cfg)
    except Exception as exc:
        print(f"\u4ea4\u53c9\u5f15\u7528\u521b\u5efa\u51fa\u9519\uff08\u5df2\u8df3\u8fc7\uff09: {exc}", file=sys.stderr)

    tbl_cfg = cfg["table"]
    tbl_cell_align = _ALIGN_MAP.get(tbl_cfg.get("cell_align", "center"))
    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:type'), 'pct')
        tblW.set(qn('w:w'), '5000')
        tblLayout = tblPr.find(qn('w:tblLayout'))
        if tblLayout is None:
            tblLayout = OxmlElement('w:tblLayout')
            tblPr.append(tblLayout)
        tblLayout.set(qn('w:type'), 'autofit')

        rows = len(table.rows)
        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    if tbl_cell_align is not None:
                        p.alignment = tbl_cell_align
                    p.paragraph_format.space_before = parse_length(0)
                    p.paragraph_format.space_after = parse_length(0)
                    p.paragraph_format.first_line_indent = parse_length(0)
                    p.paragraph_format.line_spacing = tbl_cfg["line_spacing"]
                    set_para_runs_font(p, east_asia=body_font, size_pt=caption_size,
                                       bold=False, latin=latin)

                clear_table_border(cell, "left")
                clear_table_border(cell, "right")
                clear_table_border(cell, "insideV")
                clear_table_border(cell, "insideH")

                if r_idx == 0:
                    set_table_border(cell, "top", sz=tbl_cfg["top_border_sz"])
                    set_table_border(cell, "bottom", sz=tbl_cfg["header_border_sz"])
                if r_idx == rows - 1:
                    set_table_border(cell, "bottom", sz=tbl_cfg["bottom_border_sz"])
    if cfg.get("toc", {}).get("enabled", True):
        insert_toc(doc, cfg)
    toc_match = _find_special_display(cfg, "\u76ee\u5f55", raw=True)
    first_body_h1 = find_first_body_heading(doc, cfg)
    body_started = first_body_h1 is None
    for para in doc.paragraphs:
        if get_paragraph_heading_level(para) != 1:
            continue
        if para is first_body_h1:
            body_started = True
            continue
        if not body_started:
            continue
        t = para.text.strip().replace(" ", "").replace("\u3000", "")
        if t == toc_match:
            continue
        para.paragraph_format.page_break_before = True

    custom_cover = cfg.get("cover", {}).get("custom_docx", "")
    use_custom_cover = bool(custom_cover and os.path.isfile(custom_cover))
    if use_custom_cover:
        insert_cover_and_declaration(doc, cfg, config_path, skip_cover=True)
    elif cfg["cover"]["enabled"] and not _has_cover(doc, cfg):
        insert_cover_and_declaration(doc, cfg, config_path)

    setup_page_numbers(doc, cfg)
    try:
        setup_headers(doc, cfg)
    except Exception as e:
        print(f"  [\u8b66\u544a] \u9875\u7709\u8bbe\u7f6e\u51fa\u9519\uff0c\u5df2\u8df3\u8fc7: {e}", file=sys.stderr)

    _restore_preserved_front_paragraphs()
    doc.save(output_path)
    if use_custom_cover:
        success, err = _insert_cover_via_vbs(output_path, custom_cover)
        if success:
            print("自定义封面已插入 (VBS)", file=sys.stderr)
        else:
            print(f"自定义封面插入失败（VBS不可用，已跳过）: {err}", file=sys.stderr)
    patch_theme_fonts(output_path, cfg)
    if renum_changes:
        warnings.append("\u6807\u9898\u7f16\u53f7\u5df2\u81ea\u52a8\u4fee\u6b63:")
        warnings.extend(renum_changes)
    return warnings


def patch_theme_fonts(docx_path, cfg):
    import xml.etree.ElementTree as ET
    theme = cfg.get("theme_fonts", {})
    theme_latin = theme.get("latin", "Times New Roman")
    theme_hans = theme.get("hans", "\u5b8b\u4f53")

    a_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    ET.register_namespace("a", a_ns)
    ns = {"a": a_ns}

    fd, tmp_path = tempfile.mkstemp(suffix=".docx", dir=os.path.dirname(docx_path))
    os.close(fd)
    try:
        with zipfile.ZipFile(docx_path, "r") as zin, \
             zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/theme/theme1.xml":
                    root = ET.fromstring(data)
                    for minor in root.findall(".//a:minorFont", ns):
                        lat = minor.find("a:latin", ns)
                        if lat is not None:
                            lat.set("typeface", theme_latin)
                        for font in minor.findall("a:font", ns):
                            if font.get("script") == "Hans":
                                font.set("typeface", theme_hans)
                    for major in root.findall(".//a:majorFont", ns):
                        lat = major.find("a:latin", ns)
                        if lat is not None:
                            lat.set("typeface", theme_latin)
                        for font in major.findall("a:font", ns):
                            if font.get("script") == "Hans":
                                font.set("typeface", theme_hans)
                    data = ET.tostring(root, encoding="unicode").encode("utf-8")
                zout.writestr(item, data)
        os.replace(tmp_path, docx_path)
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def main():
    import argparse
    parser = argparse.ArgumentParser(description="Universal thesis formatter")
    parser.add_argument("--input", required=True, help="Input docx")
    parser.add_argument("--output", required=True, help="Output docx")
    parser.add_argument("--config", help="Path to thesis_config.yaml")
    args = parser.parse_args()

    cfg, cfg_path = resolve_config(cli_config=args.config, input_path=args.input)
    apply_format(args.input, args.output, config=cfg, config_path=cfg_path)
    print(f"OK {args.output}")





