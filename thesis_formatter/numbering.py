"""Word 原生编号功能：多级列表 + 题注

替代文本替换方式，使用 Word 的 numbering 和 SEQ 域实现自动编号。
"""

import copy
import re

from docx import Document as DocxDocument
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.parts.numbering import NumberingPart
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from ._common import get_paragraph_heading_level

CAPTION_MODE_STABLE = "stable"
CAPTION_MODE_DYNAMIC = "dynamic"


def _make_field_runs(instr, display, rPr_el=None, font=None, size=None, latin_font=None):
    """创建 Word 域代码的 XML 元素.

    Args:
        instr: 域指令 (如 "SEQ Figure \\* ARABIC")
        display: 显示文本（仅用于需要预设显示值的域，SEQ域通常留空）
        rPr_el: 源 run 属性
        font: 字体名称 (用于覆盖)
        size: 字号 (磅)
        latin_font: 拉丁字体名称
    """
    els = []
    for ftype in ('begin', None, 'separate', None, 'end'):
        r = OxmlElement('w:r')

        # 设置 run 属性
        if rPr_el is not None:
            if font is not None or size is not None:
                # 需要修改字体，创建新的 rPr
                r_pr = copy.deepcopy(rPr_el)
                if font is not None:
                    r_fonts = r_pr.find(qn("w:rFonts"))
                    if r_fonts is None:
                        r_fonts = OxmlElement("w:rFonts")
                        r_pr.append(r_fonts)
                    r_fonts.set(qn("w:eastAsia"), font)
                    if latin_font:
                        r_fonts.set(qn("w:ascii"), latin_font)
                        r_fonts.set(qn("w:hAnsi"), latin_font)
                        r_fonts.set(qn("w:cs"), latin_font)
                if size is not None:
                    sz = r_pr.find(qn("w:sz"))
                    if sz is None:
                        sz = OxmlElement("w:sz")
                        r_pr.append(sz)
                    sz.set(qn("w:val"), str(int(size * 2)))
                    sz_cs = r_pr.find(qn("w:szCs"))
                    if sz_cs is None:
                        sz_cs = OxmlElement("w:szCs")
                        r_pr.append(sz_cs)
                    sz_cs.set(qn("w:val"), str(int(size * 2)))
                r.append(r_pr)
            else:
                r.append(copy.deepcopy(rPr_el))

        if ftype in ('begin', 'separate', 'end'):
            fc = OxmlElement('w:fldChar')
            fc.set(qn('w:fldCharType'), ftype)
            # 标记域需要更新（仅对 end 类型设置）
            if ftype == 'end':
                fc.set(qn('w:dirty'), '1')
            r.append(fc)
        elif len(els) == 1:
            it = OxmlElement('w:instrText')
            it.set(qn('xml:space'), 'preserve')
            it.text = f' {instr} '
            r.append(it)
        else:
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = display
            r.append(t)
        els.append(r)
    return els


def get_caption_mode(cfg):
    """Return normalized caption mode."""
    mode = str(cfg.get("captions", {}).get("mode", CAPTION_MODE_STABLE)).strip().lower()
    return CAPTION_MODE_DYNAMIC if mode == CAPTION_MODE_DYNAMIC else CAPTION_MODE_STABLE


def _shorten_text(text, limit=24):
    text = text.strip()
    if len(text) <= limit:
        return text
    return text[:limit - 1] + "…"


def _find_numpr(para):
    """Find numbering properties on paragraph or style."""
    p_pr = para._element.find(qn("w:pPr"))
    if p_pr is not None:
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is not None:
            return num_pr

    style = getattr(para, "style", None)
    if style is not None:
        style_ppr = style.element.find(qn("w:pPr"))
        if style_ppr is not None:
            num_pr = style_ppr.find(qn("w:numPr"))
            if num_pr is not None:
                return num_pr
    return None


def _append_literal_run(p_el, r_pr, text, east_asia_font, size_pt, latin_font):
    """Append a literal text run with caption formatting."""
    if not text:
        return
    r = OxmlElement("w:r")
    r_pr_copy = _create_caption_rpr(r_pr, east_asia_font, size_pt, latin_font)
    if r_pr_copy is not None:
        r.append(r_pr_copy)
    t_el = OxmlElement("w:t")
    t_el.set(qn("xml:space"), "preserve")
    t_el.text = text
    r.append(t_el)
    p_el.append(r)


def precheck_dynamic_caption_mode(doc, cfg):
    """Validate whether strict dynamic caption fields are safe to use."""
    cap_cfg = cfg.get("captions", {})
    sec_cfg = cfg.get("sections", {})

    chapter_pat = re.compile(sec_cfg.get("chapter_pattern", r"^第\s*\d+\s*章\b"))
    appendix_pat = re.compile(sec_cfg.get("appendix_pattern", r"^附录\s*[A-Z]"))
    h2_pat = re.compile(sec_cfg.get("h2_pattern", r"^\d+\.\d+(\s|(?=[\u4e00-\u9fff]))"))
    h3_pat = re.compile(sec_cfg.get("h3_pattern", r"^\d+\.\d+\.\d+(\s|(?=[\u4e00-\u9fff]))"))
    h4_pat = re.compile(sec_cfg.get("h4_pattern", r"^\d+\.\d+\.\d+\.\d+(\s|(?=[\u4e00-\u9fff]))"))
    fig_pat = re.compile(cap_cfg.get("figure_pattern", r"^图\s*\d"))
    tbl_pat = re.compile(cap_cfg.get("table_pattern", r"^(续)?表\s*\d"))

    try:
        from ._titles import _get_special_title_map
        st_map = _get_special_title_map(cfg)
    except Exception:
        st_map = {}

    special_set = set(st_map.keys())
    special_set.update(
        s.replace(" ", "").replace("\u3000", "")
        for s in sec_cfg.get("special_h1", [])
    )

    manual_patterns = {
        1: chapter_pat,
        2: h2_pat,
        3: h3_pat,
        4: h4_pat,
    }
    include_chapter = cap_cfg.get("include_chapter", False)
    chapter_heading_level = int(cap_cfg.get("chapter_heading_level", 1) or 1)
    has_referenceable_heading = False
    has_caption_candidate = False

    reasons = []
    seen = set()

    def add_reason(message):
        if message not in seen:
            reasons.append(message)
            seen.add(message)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        text_nospace = text.replace(" ", "").replace("\u3000", "")
        level = get_paragraph_heading_level(para)

        suspected_level = None
        if chapter_pat.match(text):
            suspected_level = 1
        elif h2_pat.match(text):
            suspected_level = 2
        elif h3_pat.match(text):
            suspected_level = 3
        elif h4_pat.match(text):
            suspected_level = 4

        if suspected_level and level is None and text_nospace not in special_set and not appendix_pat.match(text):
            add_reason(
                f"检测到疑似标题“{_shorten_text(text)}”未使用真正的 Heading{suspected_level} 样式"
            )
            continue

        if level is None:
            continue
        if level == 1 and (text_nospace in special_set or appendix_pat.match(text)):
            continue

        manual_pattern = manual_patterns.get(level)
        if manual_pattern and manual_pattern.match(text):
            add_reason(
                f"标题“{_shorten_text(text)}”虽然是 Heading{level}，但编号仍是手打文本，不是 Word 多级列表"
            )
            continue

        num_pr = _find_numpr(para)
        if num_pr is None:
            add_reason(
                f"标题“{_shorten_text(text)}”是 Heading{level}，但未检测到 Word 多级列表编号"
            )
            continue

        if include_chapter and level == chapter_heading_level:
            has_referenceable_heading = True

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if fig_pat.match(text):
            has_caption_candidate = True
            if not _match_caption_parts(text, "图"):
                add_reason(f"图题“{_shorten_text(text)}”不是标准单段格式")
            elif any("\n" in (run.text or "") or "\v" in (run.text or "") for run in para.runs):
                add_reason(f"图题“{_shorten_text(text)}”包含换行，不是标准单段格式")
        elif tbl_pat.match(text):
            has_caption_candidate = True
            if not _match_caption_parts(text, "表"):
                add_reason(f"表题“{_shorten_text(text)}”不是标准单段格式")
            elif any("\n" in (run.text or "") or "\v" in (run.text or "") for run in para.runs):
                add_reason(f"表题“{_shorten_text(text)}”包含换行，不是标准单段格式")

    if include_chapter and has_caption_candidate and not has_referenceable_heading:
        add_reason(
            f"未检测到可供动态题注引用的 Heading{chapter_heading_level} 多级编号标题"
        )

    return reasons


def resolve_caption_mode(doc, cfg):
    """Resolve caption mode against the raw, unnormalized input document."""
    requested_mode = get_caption_mode(cfg)
    if requested_mode != CAPTION_MODE_DYNAMIC:
        return requested_mode, CAPTION_MODE_STABLE, [], []

    reasons = precheck_dynamic_caption_mode(doc, cfg)
    if reasons:
        return requested_mode, CAPTION_MODE_STABLE, [], reasons

    return requested_mode, CAPTION_MODE_DYNAMIC, [
        "  提示: dynamic 题注模式原始预检通过，将使用严格动态题注域。"
    ], []



def resolve_caption_mode_after_normalization(doc, cfg, raw_reasons=None):
    """Resolve caption mode after auto heading normalization and numbering setup."""
    requested_mode = get_caption_mode(cfg)
    if requested_mode != CAPTION_MODE_DYNAMIC:
        return requested_mode, CAPTION_MODE_STABLE, []

    raw_reasons = list(raw_reasons or [])
    normalized_reasons = precheck_dynamic_caption_mode(doc, cfg)

    if not normalized_reasons:
        if raw_reasons:
            warnings = [
                "  提示: 原始文档未通过 dynamic 题注预检，但自动补 Heading/多级编号后已通过，将使用严格动态题注域。"
            ]
            warnings.extend(f"  - 原始预检: {reason}" for reason in raw_reasons)
            return requested_mode, CAPTION_MODE_DYNAMIC, warnings
        return requested_mode, CAPTION_MODE_DYNAMIC, [
            "  提示: dynamic 题注模式预检通过，将使用严格动态题注域。"
        ]

    warnings = ["  警告: dynamic 题注模式预检未通过，已自动回退 stable。"]
    if raw_reasons:
        warnings.append("  原始文档预检问题:")
        warnings.extend(f"  - {reason}" for reason in raw_reasons)
    if not raw_reasons or normalized_reasons != raw_reasons:
        warnings.append("  自动标准化后仍存在问题:")
        warnings.extend(f"  - {reason}" for reason in normalized_reasons)
    return requested_mode, CAPTION_MODE_STABLE, warnings


def _get_or_add_numbering_part(doc):
    """Return a numbering part, cloning one from a blank document if missing."""
    try:
        return doc.part.numbering_part
    except NotImplementedError:
        blank_doc = DocxDocument()
        blank_numbering_part = blank_doc.part.numbering_part
        numbering_part = NumberingPart.load(
            blank_numbering_part.partname,
            CT.WML_NUMBERING,
            blank_numbering_part.blob,
            doc.part.package,
        )
        doc.part.relate_to(numbering_part, RT.NUMBERING)
        return numbering_part


def setup_multilevel_list(doc, cfg):
    """设置 Word 多级列表，替代文本替换方式的标题编号.

    为标题 1-4 创建多级列表，插入新标题时编号自动更新。

    Args:
        doc: python-docx Document 对象
        cfg: 配置字典

    Returns:
        list: 修改的标题列表
    """
    sec = cfg.get("sections", {})

    # 获取或创建 numbering part
    numbering_part = _get_or_add_numbering_part(doc)

    numbering_el = numbering_part.element

    # 创建新的多级列表
    abstract_num_id = _get_next_abstract_num_id(numbering_el)
    abstract_num = _create_abstract_num(abstract_num_id, cfg)
    numbering_el.append(abstract_num)

    # 创建 num 定义
    num_id = _get_next_num_id(numbering_el)
    num = _create_num(num_id, abstract_num_id)
    numbering_el.append(num)

    # 应用多级列表到标题
    changes = _apply_numbering_to_headings(doc, num_id, cfg)

    return changes


def _get_next_abstract_num_id(numbering_el):
    """获取下一个可用的 abstractNumId."""
    existing_ids = set()
    for an in numbering_el.findall(qn("w:abstractNum")):
        id_val = an.get(qn("w:abstractNumId"))
        if id_val:
            existing_ids.add(int(id_val))
    for i in range(1, 1000):
        if i not in existing_ids:
            return i
    return 1


def _get_next_num_id(numbering_el):
    """获取下一个可用的 numId."""
    existing_ids = set()
    for n in numbering_el.findall(qn("w:num")):
        id_val = n.get(qn("w:numId"))
        if id_val:
            existing_ids.add(int(id_val))
    for i in range(1, 1000):
        if i not in existing_ids:
            return i
    return 1


def _create_abstract_num(abstract_num_id, cfg):
    """创建抽象编号定义."""
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_num_id))

    # 多级列表模板
    levels = [
        {"lvl": 0, "fmt": "decimal", "txt": "%1 ", "style": "Heading1"},
        {"lvl": 1, "fmt": "decimal", "txt": "%1.%2 ", "style": "Heading2"},
        {"lvl": 2, "fmt": "decimal", "txt": "%1.%2.%3 ", "style": "Heading3"},
        {"lvl": 3, "fmt": "decimal", "txt": "%1.%2.%3.%4 ", "style": "Heading4"},
    ]

    for level_info in levels:
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level_info["lvl"]))

        # 起始值
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)

        # 编号格式
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), level_info["fmt"])
        lvl.append(num_fmt)

        # 级别文本
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), level_info["txt"])
        lvl.append(lvl_text)

        # 编号后缀：使用一个普通空格；配合 lvlText 中的一个空格，总计两个普通空格
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "space")
        lvl.append(suff)

        # 关联到段落样式
        p_style = OxmlElement("w:pStyle")
        p_style.set(qn("w:val"), level_info["style"])
        lvl.append(p_style)

        abstract_num.append(lvl)

    return abstract_num


def _create_num(num_id, abstract_num_id):
    """创建 num 定义."""
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))

    abstract_num_ref = OxmlElement("w:abstractNumId")
    abstract_num_ref.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract_num_ref)

    return num


def _apply_numbering_to_headings(doc, num_id, cfg):
    """应用多级列表到标题段落."""
    sec = cfg.get("sections", {})

    # 获取特殊标题映射
    try:
        from ._titles import _get_special_title_map
        st_map = _get_special_title_map(cfg)
    except:
        st_map = {}

    chap_pat = re.compile(sec.get("chapter_pattern", r"^第\s*\d+\s*章\b"))
    appendix_pat = re.compile(sec.get("appendix_pattern", r"^附录\s*[A-Z]"))
    h2_pat = re.compile(sec.get("h2_pattern", r"^\d+\.\d+\s"))
    h3_pat = re.compile(sec.get("h3_pattern", r"^\d+\.\d+\.\d+\s"))
    h4_pat = re.compile(sec.get("h4_pattern", r"^\d+\.\d+\.\d+\.\d+\s"))

    special_set = set(st_map.keys())
    special_set.update(s.replace(" ", "").replace("\u3000", "")
                       for s in sec.get("special_h1", []))

    changes = []
    in_appendix = False

    for para in doc.paragraphs:
        level = get_paragraph_heading_level(para)
        t = para.text.strip()
        t_nospace = t.replace(" ", "").replace("\u3000", "")
        if not t:
            continue

        # 跳过特殊标题
        if level == 1 and t_nospace in special_set:
            continue

        # 处理附录
        if level == 1 and appendix_pat.match(t):
            in_appendix = True
            continue

        # 附录内的标题不处理
        if in_appendix:
            continue

        p_el = para._element
        p_pr = p_el.find(qn("w:pPr"))
        if p_pr is None:
            p_pr = OxmlElement("w:pPr")
            p_el.insert(0, p_pr)

        # 应用多级列表
        if level == 1 and chap_pat.match(t):
            _set_numbering(p_pr, num_id, 0)
            new_t = re.sub(r'^第\s*\d+\s*章\s*', "", t).strip()
            if new_t != t:
                changes.append(f"H1: \"{t}\" → \"{new_t}\"")
                _set_para_text(para, new_t)
        elif level == 2 and h2_pat.match(t):
            _set_numbering(p_pr, num_id, 1)
            new_t = re.sub(r'^\d+\.\d+\s*', "", t).strip()
            if new_t != t:
                changes.append(f"H2: \"{t}\" → \"{new_t}\"")
                _set_para_text(para, new_t)
        elif level == 3 and h3_pat.match(t):
            _set_numbering(p_pr, num_id, 2)
            new_t = re.sub(r'^\d+\.\d+\.\d+\s*', "", t).strip()
            if new_t != t:
                changes.append(f"H3: \"{t}\" → \"{new_t}\"")
                _set_para_text(para, new_t)
        elif level == 4 and h4_pat.match(t):
            _set_numbering(p_pr, num_id, 3)
            new_t = re.sub(r'^\d+\.\d+\.\d+\.\d+\s*', "", t).strip()
            if new_t != t:
                changes.append(f"H4: \"{t}\" → \"{new_t}\"")
                _set_para_text(para, new_t)

    return changes


def _set_numbering(p_pr, num_id, lvl):
    """为段落设置编号."""
    # 移除旧的编号
    for old in p_pr.findall(qn("w:numPr")):
        p_pr.remove(old)

    num_pr = OxmlElement("w:numPr")

    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(lvl))
    num_pr.append(ilvl)

    num_id_ref = OxmlElement("w:numId")
    num_id_ref.set(qn("w:val"), str(num_id))
    num_pr.append(num_id_ref)

    p_pr.append(num_pr)


def _set_para_text(para, new_text):
    """设置段落文本，保留格式."""
    if not para.runs:
        para.text = new_text
        return

    # 保存第一个 run 的格式
    first_run = para.runs[0]
    r_pr = first_run._element.find(qn("w:rPr"))

    # 清空段落内容（保留 pPr）
    p_el = para._element
    for child in list(p_el):
        if child.tag != qn("w:pPr"):
            p_el.remove(child)

    # 添加新 run
    new_run = OxmlElement("w:r")
    if r_pr is not None:
        new_run.append(copy.deepcopy(r_pr))

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = new_text
    new_run.append(t)

    p_el.append(new_run)


def _auto_apply_heading_styles(doc, cfg):
    """自动识别章节标题并应用 Heading 样式.

    如果文档没有使用标题样式，这个函数会：
    1. 识别"第X章"、"X.X"等格式的标题
    2. 自动应用对应的 Heading 样式
    3. 使 STYLEREF 域能正常工作
    """
    sec_cfg = cfg.get("sections", {})
    chapter_pat = re.compile(sec_cfg.get("chapter_pattern", r"^第\s*\d+\s*章\b"))
    h2_pat = re.compile(sec_cfg.get("h2_pattern", r"^\d+\.\d+\s"))
    h3_pat = re.compile(sec_cfg.get("h3_pattern", r"^\d+\.\d+\.\d+\s"))

    changes = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 检查当前样式
        current_level = get_paragraph_heading_level(para)

        # 一级标题（章）
        if chapter_pat.match(text) and current_level != 1:
            para.style = doc.styles["Heading1"]
            changes.append(f'自动应用 Heading1: "{text[:30]}"')

        # 二级标题
        elif h2_pat.match(text) and current_level != 2:
            para.style = doc.styles["Heading2"]
            changes.append(f'自动应用 Heading2: "{text[:30]}"')

        # 三级标题
        elif h3_pat.match(text) and current_level != 3:
            para.style = doc.styles["Heading3"]
            changes.append(f'自动应用 Heading3: "{text[:30]}"')

    return changes


def setup_figure_captions(doc, cfg):
    """设置图题 SEQ 域.

    使用 SEQ 域替代文本编号，新增图时编号自动更新。
    """
    # 如果需要章节号，先自动识别并应用标题样式
    if cfg.get("captions", {}).get("include_chapter", False):
        _auto_apply_heading_styles(doc, cfg)

    cap_cfg = cfg.get("captions", {})
    fig_pat = cap_cfg.get("figure_pattern", r"^图\s*\d")

    changes = _apply_caption_seq(doc, fig_pat, "Figure", "图", cfg)

    return changes


def setup_table_captions(doc, cfg):
    """设置表题 SEQ 域.

    使用 SEQ 域替代文本编号，新增表时编号自动更新。
    """
    # 如果需要章节号，先自动识别并应用标题样式
    if cfg.get("captions", {}).get("include_chapter", False):
        _auto_apply_heading_styles(doc, cfg)

    cap_cfg = cfg.get("captions", {})
    tbl_pat = cap_cfg.get("table_pattern", r"^(续)?表\s*\d")

    changes = _apply_caption_seq(doc, tbl_pat, "Table", "表", cfg)

    return changes


def _apply_caption_seq(doc, pattern, seq_name, label, cfg):
    """应用 SEQ 域到图/表题，按章编号时保留稳定的静态前缀编号."""
    cap_cfg = cfg.get("captions", {})
    pat = re.compile(pattern)
    changes = []

    # 题注字体配置
    cap_font = cap_cfg.get("font", "宋体")
    num_font = cap_cfg.get("number_font", "Times New Roman")
    cap_size = cap_cfg.get("size", 10.5)
    latin_font = cfg.get("fonts", {}).get("latin", "Times New Roman")

    # 编号格式配置
    include_chapter = cap_cfg.get("include_chapter", False)
    chapter_sep = cap_cfg.get("chapter_separator", ".")
    caption_sep = cap_cfg.get("caption_separator", "")
    chapter_heading_level = cap_cfg.get("chapter_heading_level", 1)
    restart_per_chapter = cap_cfg.get("restart_per_chapter", False)

    for para in doc.paragraphs:
        t = para.text.strip()
        if not pat.match(t):
            continue

        parts = _match_caption_parts(t, label)
        if not parts:
            continue

        prefix = parts.group("prefix").rstrip()
        normalized_number = re.sub(r"\s*([.\-])\s*", r"\1", parts.group("number").strip())
        suffix_raw = parts.group("suffix") or ""
        suffix_text = suffix_raw.lstrip() if caption_sep else suffix_raw
        number_tokens = [token for token in re.split(r"[.\-]", normalized_number) if token]

        def token_or_empty(index):
            return number_tokens[index] if index < len(number_tokens) else ""

        use_chapter_number = include_chapter
        reset_with_chapter = restart_per_chapter
        effective_chapter_sep = chapter_sep

        p_el = para._element
        p_pr = p_el.find(qn("w:pPr"))
        if p_pr is None:
            p_pr = OxmlElement("w:pPr")
            p_el.insert(0, p_pr)

        r_pr = None
        if para.runs:
            r_pr = para.runs[0]._element.find(qn("w:rPr"))

        for child in list(p_el):
            if child.tag != qn("w:pPr"):
                p_el.remove(child)

        r = OxmlElement("w:r")
        r_pr_copy = _create_caption_rpr(r_pr, cap_font, cap_size, latin_font)
        if r_pr_copy is not None:
            r.append(r_pr_copy)
        t_el = OxmlElement("w:t")
        t_el.set(qn("xml:space"), "preserve")
        t_el.text = prefix
        r.append(t_el)
        p_el.append(r)

        if use_chapter_number:
            full_number = normalized_number
            r = OxmlElement("w:r")
            r_pr_copy = _create_caption_rpr(r_pr, num_font, cap_size, latin_font)
            if r_pr_copy is not None:
                r.append(r_pr_copy)
            t_el = OxmlElement("w:t")
            t_el.set(qn("xml:space"), "preserve")
            t_el.text = full_number
            r.append(t_el)
            p_el.append(r)
        else:
            seq_display = token_or_empty(-1)
            seq_instr = f"SEQ {seq_name} \* ARABIC"
            for fel in _make_field_runs(seq_instr, seq_display, r_pr, num_font, cap_size, latin_font):
                p_el.append(fel)

        if caption_sep:
            r = OxmlElement("w:r")
            r_pr_copy = _create_caption_rpr(r_pr, num_font, cap_size, latin_font)
            if r_pr_copy is not None:
                r.append(r_pr_copy)
            t_el = OxmlElement("w:t")
            t_el.set(qn("xml:space"), "preserve")
            t_el.text = caption_sep
            r.append(t_el)
            p_el.append(r)

        if suffix_text:
            r = OxmlElement("w:r")
            r_pr_copy = _create_caption_rpr(r_pr, cap_font, cap_size, latin_font)
            if r_pr_copy is not None:
                r.append(r_pr_copy)
            t_el = OxmlElement("w:t")
            t_el.set(qn("xml:space"), "preserve")
            t_el.text = suffix_text
            r.append(t_el)
            p_el.append(r)

        changes.append(f'{label}题注域化: "{t}"')

    return changes


def _apply_caption_dynamic(doc, pattern, seq_name, label, cfg):
    """Apply strict dynamic STYLEREF + SEQ fields to captions."""
    cap_cfg = cfg.get("captions", {})
    pat = re.compile(pattern)
    changes = []

    cap_font = cap_cfg.get("font", "宋体")
    num_font = cap_cfg.get("number_font", "Times New Roman")
    cap_size = cap_cfg.get("size", 10.5)
    latin_font = cfg.get("fonts", {}).get("latin", "Times New Roman")

    include_chapter = cap_cfg.get("include_chapter", False)
    chapter_sep = cap_cfg.get("chapter_separator", ".")
    caption_sep = cap_cfg.get("caption_separator", "")
    chapter_heading_level = cap_cfg.get("chapter_heading_level", 1)
    restart_per_chapter = cap_cfg.get("restart_per_chapter", False)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not pat.match(text):
            continue

        parts = _match_caption_parts(text, label)
        if not parts:
            continue

        prefix = parts.group("prefix").rstrip()
        normalized_number = re.sub(r"\s*([.\-])\s*", r"\1", parts.group("number").strip())
        suffix_raw = parts.group("suffix") or ""
        suffix_text = suffix_raw.lstrip() if caption_sep else suffix_raw
        number_tokens = [token for token in re.split(r"[.\-]", normalized_number) if token]

        use_chapter_number = include_chapter and len(number_tokens) > 1
        chapter_display = chapter_sep.join(number_tokens[:-1]) if use_chapter_number else ""
        seq_display = number_tokens[-1] if number_tokens else ""
        is_continued_table = label == "表" and prefix.startswith("续")

        p_el = para._element
        p_pr = p_el.find(qn("w:pPr"))
        if p_pr is None:
            p_pr = OxmlElement("w:pPr")
            p_el.insert(0, p_pr)

        r_pr = None
        if para.runs:
            r_pr = para.runs[0]._element.find(qn("w:rPr"))

        for child in list(p_el):
            if child.tag != qn("w:pPr"):
                p_el.remove(child)

        _append_literal_run(p_el, r_pr, prefix, cap_font, cap_size, latin_font)

        if use_chapter_number:
            for fel in _make_styleref_field(
                r_pr, chapter_heading_level, num_font, cap_size, latin_font, chapter_display
            ):
                p_el.append(fel)
            _append_literal_run(p_el, r_pr, chapter_sep, num_font, cap_size, latin_font)

        seq_instr_parts = [f"SEQ {seq_name}"]
        if is_continued_table:
            seq_instr_parts.append(r"\c")
        else:
            seq_instr_parts.append(r"\* ARABIC")
            if use_chapter_number and restart_per_chapter:
                seq_instr_parts.append(rf"\s {chapter_heading_level}")
        seq_instr = " ".join(seq_instr_parts)
        for fel in _make_field_runs(seq_instr, seq_display, r_pr, num_font, cap_size, latin_font):
            p_el.append(fel)

        if caption_sep:
            _append_literal_run(p_el, r_pr, caption_sep, num_font, cap_size, latin_font)
        if suffix_text:
            _append_literal_run(p_el, r_pr, suffix_text, cap_font, cap_size, latin_font)

        changes.append(f'{label}题注动态域化: "{text}"')

    return changes


def _make_styleref_field(rPr_el, heading_level, font, size, latin_font, display_text=""):
    """创建 STYLEREF 域元素，用于引用标题编号."""
    els = []
    for ftype in ('begin', None, 'separate', None, 'end'):
        r = OxmlElement('w:r')

        # 设置 run 属性
        if rPr_el is not None:
            r_pr = copy.deepcopy(rPr_el)
            if font is not None:
                r_fonts = r_pr.find(qn("w:rFonts"))
                if r_fonts is None:
                    r_fonts = OxmlElement("w:rFonts")
                    r_pr.append(r_fonts)
                r_fonts.set(qn("w:eastAsia"), font)
                if latin_font:
                    r_fonts.set(qn("w:ascii"), latin_font)
                    r_fonts.set(qn("w:hAnsi"), latin_font)
                    r_fonts.set(qn("w:cs"), latin_font)
            if size is not None:
                sz = r_pr.find(qn("w:sz"))
                if sz is None:
                    sz = OxmlElement("w:sz")
                    r_pr.append(sz)
                sz.set(qn("w:val"), str(int(size * 2)))
                sz_cs = r_pr.find(qn("w:szCs"))
                if sz_cs is None:
                    sz_cs = OxmlElement("w:szCs")
                    r_pr.append(sz_cs)
                sz_cs.set(qn("w:val"), str(int(size * 2)))
            r.append(r_pr)
        else:
            r_pr = OxmlElement("w:rPr")
            if font is not None:
                r_fonts = OxmlElement("w:rFonts")
                r_fonts.set(qn("w:eastAsia"), font)
                if latin_font:
                    r_fonts.set(qn("w:ascii"), latin_font)
                    r_fonts.set(qn("w:hAnsi"), latin_font)
                    r_fonts.set(qn("w:cs"), latin_font)
                r_pr.append(r_fonts)
            if size is not None:
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), str(int(size * 2)))
                r_pr.append(sz)
                sz_cs = OxmlElement("w:szCs")
                sz_cs.set(qn("w:val"), str(int(size * 2)))
                r_pr.append(sz_cs)
            r.append(r_pr)

        if ftype in ('begin', 'separate', 'end'):
            fc = OxmlElement('w:fldChar')
            fc.set(qn('w:fldCharType'), ftype)
            if ftype == 'end':
                fc.set(qn('w:dirty'), '1')
            r.append(fc)
        elif len(els) == 1:
            it = OxmlElement('w:instrText')
            it.set(qn('xml:space'), 'preserve')
            it.text = f' STYLEREF {heading_level} \s '
            r.append(it)
        else:
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = display_text
            r.append(t)
        els.append(r)
    return els


def _match_caption_parts(text, label):
    """提取题注前缀、编号和后缀文本."""
    label_pat = r"(?:续)?表" if label == "表" else re.escape(label)
    return re.match(
        rf"^(?P<prefix>{label_pat}\s*)(?P<number>[A-Z]?\d+(?:\s*[.\-]\s*\d+)*)(?P<suffix>\s*.*)$",
        text,
        re.I,
    )


def setup_figure_captions(doc, cfg):
    """设置图题域，按运行时模式选择 stable / dynamic。"""
    cap_cfg = cfg.get("captions", {})
    fig_pat = cap_cfg.get("figure_pattern", r"^图\s*\d")
    mode = cfg.get("_runtime", {}).get("caption_mode_effective", get_caption_mode(cfg))
    apply_func = _apply_caption_dynamic if mode == CAPTION_MODE_DYNAMIC else _apply_caption_seq

    changes = apply_func(doc, fig_pat, "Figure", "图", cfg)

    return changes


def setup_table_captions(doc, cfg):
    """设置表题域，按运行时模式选择 stable / dynamic。"""
    cap_cfg = cfg.get("captions", {})
    tbl_pat = cap_cfg.get("table_pattern", r"^(续)?表\s*\d")
    mode = cfg.get("_runtime", {}).get("caption_mode_effective", get_caption_mode(cfg))
    apply_func = _apply_caption_dynamic if mode == CAPTION_MODE_DYNAMIC else _apply_caption_seq

    changes = apply_func(doc, tbl_pat, "Table", "表", cfg)

    return changes

def _create_caption_rpr(source_rpr, east_asia_font, size_pt, latin_font):
    """创建题注的 run 属性，设置字体."""
    if source_rpr is not None:
        r_pr = copy.deepcopy(source_rpr)
    else:
        r_pr = OxmlElement("w:rPr")

    # 设置字体
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)

    r_fonts.set(qn("w:eastAsia"), east_asia_font)
    r_fonts.set(qn("w:ascii"), latin_font)
    r_fonts.set(qn("w:hAnsi"), latin_font)
    r_fonts.set(qn("w:cs"), latin_font)

    # 设置字号
    sz = r_pr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        r_pr.append(sz)
    sz.set(qn("w:val"), str(int(size_pt * 2)))

    sz_cs = r_pr.find(qn("w:szCs"))
    if sz_cs is None:
        sz_cs = OxmlElement("w:szCs")
        r_pr.append(sz_cs)
    sz_cs.set(qn("w:val"), str(int(size_pt * 2)))

    return r_pr

