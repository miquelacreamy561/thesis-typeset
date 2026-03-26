import copy
import re

from ._common import _ALIGN_MAP, set_rfonts, get_paragraph_heading_level, parse_length
from ._titles import _find_special_display
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


from docx.shared import Cm


_AR_CHAPTER_PATTERN = r"^第\s*\d+\s*章(?:\s|(?=[\u4e00-\u9fff])|$)"
_CN_CHAPTER_PATTERN = r"^第\s*[一二三四五六七八九十百千零两〇]+\s*章(?:\s|(?=[\u4e00-\u9fff])|$)"


def _iter_chapter_patterns(sec, text_first=False):
    configured = sec.get("chapter_pattern", _AR_CHAPTER_PATTERN)
    ordered = [_CN_CHAPTER_PATTERN, _AR_CHAPTER_PATTERN, configured] if text_first else [configured, _CN_CHAPTER_PATTERN, _AR_CHAPTER_PATTERN]
    seen = set()
    for pat in ordered:
        if pat and pat not in seen:
            seen.add(pat)
            yield pat


def _matches_chapter_heading(text, sec, text_first=False):
    return any(re.match(pat, text) for pat in _iter_chapter_patterns(sec, text_first=text_first))


def normalize_sections(doc, cfg):
    pg = cfg["page"]
    for section in doc.sections:
        section.top_margin = Cm(pg["margins"]["top"])
        section.bottom_margin = Cm(pg["margins"]["bottom"])
        section.left_margin = Cm(pg["margins"]["left"])
        section.right_margin = Cm(pg["margins"]["right"])
        section.gutter = Cm(pg["gutter"])
        section.header_distance = Cm(pg["header_distance"])
        section.footer_distance = Cm(pg["footer_distance"])
        try:
            for p in section.header.paragraphs:
                p.clear()
        except Exception:
            pass


def add_page_number_field(paragraph, cfg, align="center"):
    paragraph.alignment = _ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    pn_size = parse_length(cfg["sizes"]["page_number"])
    pn_cfg = cfg["page_numbers"]
    pn_font = pn_cfg.get("font", "") or cfg["fonts"]["latin"]
    pn_bold = pn_cfg.get("bold", False)
    body_ea = cfg["fonts"]["body"]
    decorator = pn_cfg.get("decorator", "{page}")
    prefix, suffix = "", ""
    if "{page}" in decorator:
        parts = decorator.split("{page}", 1)
        prefix, suffix = parts[0], parts[1]
    elif decorator != "{page}":
        prefix = ""

    def _pn_run(text=None):
        r = paragraph.add_run(text) if text else paragraph.add_run()
        r.font.size = pn_size
        r.font.name = pn_font
        r.font.bold = pn_bold
        rpr = r._element.get_or_add_rPr()
        set_rfonts(rpr, body_ea, pn_font)
        return r

    if prefix:
        _pn_run(prefix)

    run = _pn_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin)

    run2 = _pn_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._element.append(instr)

    run3 = _pn_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._element.append(fld_sep)

    _pn_run("1")

    run5 = _pn_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run5._element.append(fld_char_end)

    if suffix:
        _pn_run(suffix)


def set_section_page_number_format(section, fmt="decimal", start=None):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), fmt)
    start_key = qn("w:start")
    if start is not None:
        pg_num.set(start_key, str(start))
    elif start_key in pg_num.attrib:
        del pg_num.attrib[start_key]


def _is_page_break_only_paragraph(p_element):
    if p_element is None or p_element.tag != qn("w:p"):
        return False
    if "".join(p_element.itertext()).strip():
        return False
    for br in p_element.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def insert_section_break_before(paragraph):
    p_element = paragraph._element
    prev = p_element.getprevious()
    if prev is None:
        return None

    break_holder = prev
    if _is_page_break_only_paragraph(prev):
        prev_prev = prev.getprevious()
        if prev_prev is not None and prev_prev.tag == qn("w:p"):
            break_holder = prev_prev
            parent = prev.getparent()
            if parent is not None:
                parent.remove(prev)
        else:
            break_holder = prev
            pPr = break_holder.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                break_holder.insert(0, pPr)
            for child in list(break_holder):
                if child is not pPr:
                    break_holder.remove(child)

    if break_holder.tag != qn("w:p"):
        break_holder = OxmlElement("w:p")
        p_element.addprevious(break_holder)

    pPr = break_holder.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        break_holder.insert(0, pPr)

    for existing in list(pPr.findall(qn("w:sectPr"))):
        pPr.remove(existing)

    sect_pr = OxmlElement("w:sectPr")
    sect_type = OxmlElement("w:type")
    sect_type.set(qn("w:val"), "nextPage")
    sect_pr.append(sect_type)
    pPr.append(sect_pr)
    return sect_pr


def insert_page_break_after(paragraph):
    p_element = paragraph._element
    new_p = OxmlElement("w:p")
    new_r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    new_r.append(br)
    new_p.append(new_r)
    p_element.addnext(new_p)


def _normalize_title(text):
    return text.replace(" ", "").replace("\u3000", "")





def find_first_body_heading(doc, cfg):
    sec = cfg.get("sections", {})
    text_first = bool(cfg.get("toc", {}).get("only_insert", False))
    appendix_re = re.compile(sec.get("appendix_pattern", r"^附录\s*[A-Z]"))

    skip_titles = set()
    for entry in cfg.get("special_titles", []):
        match = entry.get("match", "")
        display = entry.get("display", "")
        if match:
            skip_titles.add(_normalize_title(match))
        if display:
            skip_titles.add(_normalize_title(display))
    for title in sec.get("special_h1", []):
        skip_titles.add(_normalize_title(title))
    skip_titles.add(_normalize_title(_find_special_display(cfg, "目录", raw=True)))

    headings = []
    for para in doc.paragraphs:
        if get_paragraph_heading_level(para) != 1:
            continue
        text = para.text.strip()
        if not text:
            continue
        headings.append((para, text, _normalize_title(text)))

    for para, text, normalized in headings:
        if normalized in skip_titles or appendix_re.match(text):
            continue
        if _matches_chapter_heading(text, sec, text_first=text_first):
            return para

    for para, text, normalized in headings:
        if normalized in skip_titles or appendix_re.match(text):
            continue
        return para

    return None


def get_body_start_section_index(doc, cfg, first_body_h1=None):
    if first_body_h1 is None:
        first_body_h1 = find_first_body_heading(doc, cfg)
    if first_body_h1 is None:
        return len(doc.sections) - 1 if len(doc.sections) > 1 else 0

    target = first_body_h1._element
    section_idx = 0
    for child in doc._element.body.iterchildren():
        if child is target:
            return min(section_idx, len(doc.sections) - 1)
        if child.tag != qn("w:p"):
            continue
        ppr = child.find(qn("w:pPr"))
        if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
            section_idx += 1
    return len(doc.sections) - 1 if len(doc.sections) > 1 else 0


def _enable_even_odd_headers(section):
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:titlePg"))
    if existing is None:
        sect_pr.append(OxmlElement("w:titlePg"))


def _set_even_odd_on_doc(doc):
    settings = doc.settings.element
    if settings.find(qn("w:evenAndOddHeaders")) is None:
        settings.append(OxmlElement("w:evenAndOddHeaders"))


def _clear_story_content(story):
    story.is_linked_to_previous = False
    for p in story.paragraphs:
        p.clear()


def _clear_section_footers(section):
    for story in (section.footer, section.first_page_footer, section.even_page_footer):
        _clear_story_content(story)


def _finalize_cover_section_page_numbers(doc, cfg, body_section_index=None):
    cover_sections = int(cfg.get("_runtime", {}).get("custom_cover_sections", 0) or 0)
    if cover_sections <= 0:
        return
    cover_sections = min(cover_sections, len(doc.sections))
    for idx in range(cover_sections):
        _clear_section_footers(doc.sections[idx])
    first_non_cover_idx = cover_sections if cover_sections < len(doc.sections) else None
    if first_non_cover_idx is None:
        return
    if body_section_index is not None and first_non_cover_idx >= body_section_index:
        return
    pn = cfg["page_numbers"]
    set_section_page_number_format(doc.sections[first_non_cover_idx], fmt=pn["front_format"], start=pn["front_start"])

def _setup_single_section_pn(doc, cfg):
    pn = cfg["page_numbers"]
    body_pos = pn.get("body_position", "center")
    body_odd = pn.get("body_odd_position", "right")
    body_even = pn.get("body_even_position", "left")
    need_alternate = body_pos == "alternate"
    hf_diff_oe = cfg.get("header_footer", {}).get("different_odd_even", False) and \
                 cfg.get("header_footer", {}).get("enabled", False)
    need_even_odd = need_alternate or hf_diff_oe

    section = doc.sections[0]
    set_section_page_number_format(section, fmt=pn["body_format"], start=pn["body_start"])

    if need_even_odd:
        _set_even_odd_on_doc(doc)

    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()

    if need_alternate:
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        add_page_number_field(fp, cfg, align=body_odd)
        even_footer = section.even_page_footer
        even_footer.is_linked_to_previous = False
        for p in even_footer.paragraphs:
            p.clear()
        ep = even_footer.paragraphs[0] if even_footer.paragraphs else even_footer.add_paragraph()
        add_page_number_field(ep, cfg, align=body_even)
    else:
        actual_pos = body_pos if body_pos != "alternate" else "center"
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        add_page_number_field(fp, cfg, align=actual_pos)
        if need_even_odd:
            even_footer = section.even_page_footer
            even_footer.is_linked_to_previous = False
            for p in even_footer.paragraphs:
                p.clear()
            ep = even_footer.paragraphs[0] if even_footer.paragraphs else even_footer.add_paragraph()
            add_page_number_field(ep, cfg, align=actual_pos)


def _apply_page_numbers_to_sections(doc, cfg, body_section_index):
    pn = cfg["page_numbers"]
    front_pos = pn.get("front_position", "center")
    body_pos = pn.get("body_position", "center")
    body_odd = pn.get("body_odd_position", "right")
    body_even = pn.get("body_even_position", "left")
    need_alternate = body_pos == "alternate"
    hf_diff_oe = cfg.get("header_footer", {}).get("different_odd_even", False) and \
                 cfg.get("header_footer", {}).get("enabled", False)
    need_even_odd = need_alternate or hf_diff_oe

    for idx, section in enumerate(doc.sections):
        if idx < body_section_index:
            start = pn["front_start"] if idx == 0 else None
            set_section_page_number_format(section, fmt=pn["front_format"], start=start)
        else:
            start = pn["body_start"] if idx == body_section_index else None
            set_section_page_number_format(section, fmt=pn["body_format"], start=start)

    if need_even_odd:
        _set_even_odd_on_doc(doc)

    for idx, section in enumerate(doc.sections):
        is_body = idx >= body_section_index
        pos = body_pos if is_body else front_pos

        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()

        if pos == "alternate" and is_body:
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            add_page_number_field(fp, cfg, align=body_odd)
            even_footer = section.even_page_footer
            even_footer.is_linked_to_previous = False
            for p in even_footer.paragraphs:
                p.clear()
            ep = even_footer.paragraphs[0] if even_footer.paragraphs else even_footer.add_paragraph()
            add_page_number_field(ep, cfg, align=body_even)
        else:
            actual_pos = pos if pos != "alternate" else "center"
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            add_page_number_field(fp, cfg, align=actual_pos)
            if need_even_odd:
                even_footer = section.even_page_footer
                even_footer.is_linked_to_previous = False
                for p in even_footer.paragraphs:
                    p.clear()
                ep = even_footer.paragraphs[0] if even_footer.paragraphs else even_footer.add_paragraph()
                add_page_number_field(ep, cfg, align=actual_pos)

    _finalize_cover_section_page_numbers(
        doc,
        cfg,
        body_section_index=body_section_index if len(doc.sections) > 1 else None,
    )


def setup_page_numbers(doc, cfg):
    first_body_h1 = find_first_body_heading(doc, cfg)
    if first_body_h1 is None:
        _setup_single_section_pn(doc, cfg)
        _finalize_cover_section_page_numbers(doc, cfg)
        return

    new_sect_pr = insert_section_break_before(first_body_h1)
    if new_sect_pr is None:
        _setup_single_section_pn(doc, cfg)
        _finalize_cover_section_page_numbers(doc, cfg)
        return

    for attr in ["pgSz", "pgMar"]:
        existing = doc.sections[0]._sectPr.find(qn(f"w:{attr}"))
        if existing is not None:
            new_sect_pr.append(copy.deepcopy(existing))

    body_section_index = get_body_start_section_index(doc, cfg, first_body_h1)
    _apply_page_numbers_to_sections(doc, cfg, body_section_index)


def setup_page_numbers_strict(doc, cfg):
    if len(doc.sections) <= 1:
        _setup_single_section_pn(doc, cfg)
        _finalize_cover_section_page_numbers(doc, cfg)
        return

    first_body_h1 = find_first_body_heading(doc, cfg)
    body_section_index = get_body_start_section_index(doc, cfg, first_body_h1)
    _apply_page_numbers_to_sections(doc, cfg, body_section_index)
